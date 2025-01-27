from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from django.utils.crypto import get_random_string
from django.core.mail import send_mail
from rest_framework.permissions import IsAuthenticated, AllowAny
from django.contrib.auth import get_user_model
from rest_framework_simplejwt.tokens import RefreshToken
from django.contrib.auth import authenticate
from clientadmin.models import Chatbot, KnowledgeBase,KnowledgeBaseDocument
from clientuser.models import Conversation
from clientadmin.serializers import ChatbotSerializer, KnowledgeBaseSerializer
# from clientuser.serializers import ConversationSerializer
from django.core.files.uploadedfile import UploadedFile
from django.shortcuts import get_object_or_404
from django.urls import reverse
from account.models import User,UserAdminUserRelationship
import boto3
from django.conf import settings
from rest_framework.exceptions import ValidationError
from django.utils import timezone
from django.db import IntegrityError, transaction
import uuid 
from langchain_community.document_loaders import PyPDFLoader
from langchain_community.document_loaders import TextLoader
from langchain.text_splitter import CharacterTextSplitter,RecursiveCharacterTextSplitter
from helper.helper import PineconeInitializer
import traceback
import PyPDF2
from io import BytesIO
import io
import base64
from django.core.files.uploadedfile import InMemoryUploadedFile
import logging
from PyPDF2 import PdfReader
from django.core.files.storage import default_storage
from docx import Document
from django.utils.crypto import get_random_string
import aiohttp
from PyPDF2.errors import PdfReadError
from superadmin.models import APISettings
from django.core.exceptions import ImproperlyConfigured

from colorama import Fore
from google_auth_oauthlib.flow import Flow
import json
from .models import GoogleDriveAccount, MicrosoftAccount
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build

from django.contrib.auth.decorators import login_required
from django.utils.decorators import method_decorator
from django.utils.http import urlencode
import requests
from django.http import HttpResponseRedirect
import jwt
from google.auth.transport.requests import Request
# from datetime import datetime, timedelta
from asgiref.sync import sync_to_async
from django.utils.decorators import method_decorator
from bs4 import BeautifulSoup
import msal
import os
from urllib.parse import urlparse


from colorama import Fore
logger = logging.getLogger(__name__)



User = get_user_model()
s3_client = boto3.client(
    's3',
    aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
    aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY,
    region_name=settings.AWS_S3_REGION_NAME
)

class CreateUserView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        user_admin = request.user

        if not user_admin.is_user_admin:
            return Response({'error': 'Only User Admins can create new users'}, status=status.HTTP_403_FORBIDDEN)

        email = request.data.get('email')

        if not email:
            return Response({'error': 'Email is required'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            # Start a database transaction
            with transaction.atomic():
                try:
                    # Attempt to retrieve the user by email
                    user = User.objects.get(email=email)
                except User.DoesNotExist:
                    # If the user does not exist, create a new user
                    user = User.objects.create(
                        username=email,
                        email=email,
                        is_active=False,
                        password_reset_token=get_random_string(20),
                    )
                    user_created = True
                else:
                    user_created = False

                # Check if the user is a User Admin or Super Admin
                if user.is_user_admin or user.is_super_admin:
                    return Response({
                        'error': 'You cannot add a User Admin or Super Admin as a regular user.'
                    }, status=status.HTTP_400_BAD_REQUEST)

                # Check if the user is already associated with this User Admin
                if UserAdminUserRelationship.objects.filter(user=user, user_admin=user_admin).exists():
                    return Response({
                        'message': 'This user is already associated with your account.',
                        'email': user.email,
                        'user_status':user.is_active
                    }, status=status.HTTP_200_OK)

                # If the user is newly created or not yet associated, create the relationship
                UserAdminUserRelationship.objects.create(user=user, user_admin=user_admin)

                if user_created:
                    # If the user was created, send the password setup email
                    setup_link = "https://ai-rag-user-git-staging-octalooptechnologies-projects.vercel.app/setup-account"  # Adjust your front-end URL here
                    send_mail(
                        'Set Up Your Account Password',
                        f'Hello,\n\nYour account has been created by an admin. Please set your password by following this link:\n\n{setup_link}\n\n your token is {user.password_reset_token} \n\nThis link will expire in 24 hours.',
                        'no-reply@example.com',
                        [user.email],
                        fail_silently=False,
                    )
                    return Response({
                        'message': 'invitation sent.',
                        'email': user.email,
                        'user_status':user.is_active,
                    }, status=status.HTTP_201_CREATED)

                else:
                    # If the user already existed but was not associated, notify the admin
                    return Response({
                        'message': 'User already existed. Access has been granted to the new User Admin\'s bots.',
                        'email': user.email,
                        'user_status': user.is_active,
                    }, status=status.HTTP_200_OK)

        except IntegrityError:
            # Handle cases where database constraints are violated
            return Response({'error': 'A user with this email or username already exists.'}, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            # Handle any other exceptions that may occur
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)



class SetPasswordView(APIView):
    def post(self, request, token):
        try:
            user = User.objects.get(password_reset_token=token)
        except User.DoesNotExist:
            return Response({'error': 'Invalid token'}, status=status.HTTP_400_BAD_REQUEST)

        # Check if the token is expired (assuming 24 hours expiry)
        token_age = timezone.now() - user.date_joined  # Assuming date_joined is when the token was created
        if token_age.total_seconds() > 86400:
            return Response({'error': 'Token has expired'}, status=status.HTTP_400_BAD_REQUEST)

        # Set the new password and username
        new_password = request.data.get('password')
        username = request.data.get('username')
        if not new_password or not username:
            return Response({'error': 'Username and password are required'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            # Check if the username already exists
            if User.objects.filter(username=username).exists():
                return Response({'error': 'Username already taken, please choose another.'}, status=status.HTTP_400_BAD_REQUEST)

            user.set_password(new_password)
            user.username = username
            user.is_active = True  # Activate the user after password is set
            user.password_reset_token = ''  # Clear the token
            user.save()

            return Response({'message': 'Password has been set successfully. You can now log in.'}, status=status.HTTP_200_OK)
        except IntegrityError:
            return Response({'error': 'An error occurred while setting the password. Please try again.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    


class ListUsersCreatedByAdminView(APIView):
    permission_classes = [IsAuthenticated]
    def get(self, request):
        user_admin = request.user

        if not user_admin.is_user_admin:
            return Response({'error': 'Only User Admins can view this list'}, status=status.HTTP_403_FORBIDDEN)

        # Get all users created by this User Admin
        users = User.objects.filter(admin_relationships__user_admin=user_admin)
        if not users:
            return Response({'result':'no users'},status=status.HTTP_200_OK)
            
        
        users_data = []
        for user in users:
            # Count the number of knowledge bases created by this user
            knowledge_base_count = KnowledgeBase.objects.filter(created_by=user).count()

            # Prepare user data with additional fields
            user_data = {
                'user_id': user.id,
                'username': user.username,
                'email': user.email,
                'profile_image': user.profile_photo_url,
                'user_status': user.is_active,
                'role': user.role,  # Assuming the role field exists in the User model
                'knowledge_base_count': knowledge_base_count  # Add the count of knowledge bases
            }
            users_data.append(user_data)

        return Response({'users': users_data}, status=status.HTTP_200_OK)


     



    

class UserChatbotAccessView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        user = request.user

        # Get all chatbots associated with all the User Admins who have created this user
        user_admin_ids = user.admin_relationships.values_list('user_admin', flat=True)
        chatbots = Chatbot.objects.filter(created_by__in=user_admin_ids).order_by('-created_at')

        if not chatbots.exists():
            return Response({'message': 'No chatbots found for this user'}, status=status.HTTP_404_NOT_FOUND)

        serializer = ChatbotSerializer(chatbots, many=True)
        return Response(serializer.data, status=status.HTTP_200_OK)


class DeleteUserView(APIView):
    permission_classes = [IsAuthenticated]

    def delete(self, request, user_id):
        user_admin = request.user

        if not user_admin.is_user_admin:
            return Response({'error': 'Only User Admins can delete users'}, status=status.HTTP_403_FORBIDDEN)

        # Find the user to be deleted
        user_to_delete = get_object_or_404(User, id=user_id)

        # Check if the current User Admin is associated with this user
        if not user_to_delete.admin_relationships.filter(user_admin=user_admin).exists():
            return Response({'error': 'You do not have permission to delete this user'}, status=status.HTTP_403_FORBIDDEN)

        # Delete the user
        user_to_delete.delete()

        return Response({'message': 'User deleted successfully'}, status=status.HTTP_200_OK)






#knowledge base k endpoints

# class CreateKnowledgeBaseView(APIView):
#     permission_classes = [IsAuthenticated]

#     def post(self, request):
#         # Retrieve basic data
#         name = request.data.get('name')
#         admin_id = request.data.get('admin_id')
#         folder_id = request.data.get('folder_id')
#         chunk_size = request.data.get('chunk_size', 100)
#         chunk_overlap = request.data.get('chunk_overlap', 20)
#         separator = request.data.get('separator', ["\n\n", "\n", " ", ""])

#         # print("checkpoint 1")

#         if not name or not admin_id:
#             return Response({'error': 'Missing required fields: name and admin_id'}, status=status.HTTP_400_BAD_REQUEST)

#         # Check permissions and get user admin
#         user = request.user
#         if user.is_user_admin:
#             user_admin = user
#             creator = user
#         else:
#             if user.role != 'editor':
#                 return Response({'error': 'Only User Admins or Editors can create knowledge bases'}, status=status.HTTP_403_FORBIDDEN)
#             try:
#                 user_admin = User.objects.get(id=admin_id, is_user_admin=True)
#             except User.DoesNotExist:
#                 return Response({'error': 'Admin not found'}, status=status.HTTP_404_NOT_FOUND)

#             # Ensure the editor is associated with the admin
#             if not UserAdminUserRelationship.objects.filter(user=user, user_admin=user_admin).exists():
#                 return Response({'error': 'You are not associated with the provided admin.'}, status=status.HTTP_403_FORBIDDEN)

#             creator = user
#         # print("checkpoint 2")
#         # Check if the admin already has a Pinecone index, create if needed
#         pinecone_initializer = PineconeInitializer(pinecone_api=settings.PINECONE_API, open_ai_api=settings.OPENAI_API_KEY)
#         if not user_admin.pinecone_index:
#             index_name = f"index-{user_admin.id}"
#             pinecone_initializer.initialize_pinecone(index_name)
#             user_admin.pinecone_index = index_name
#             user_admin.save()
#         else:
#             index_name = user_admin.pinecone_index
#         # print("checkpoint 3")
#         # Check if the namespace already exists
#         if KnowledgeBase.objects.filter(created_by=user_admin, namespace=name).exists():
#             return Response({'error': 'Namespace already exists for this admin.'}, status=status.HTTP_400_BAD_REQUEST)
#         # print("checkpoint 4")
#         # Initialize for tracking S3 URLs
#         s3_urls = []

#         documents = request.data.get('documents', [])
#         # print("checkpoint 5")
#         for index, document_data in enumerate(documents):
#             base64_file_content = document_data.get('file')  # base64 string
#             document_type = document_data.get('document_type')
#             document_name = document_data.get('document_name')

#             if not document_type or not base64_file_content:
#                 return Response({'error': f'Missing file or document_type for document at index {index}'}, status=status.HTTP_400_BAD_REQUEST)

#             # Decode base64 content to bytes
#             if base64_file_content.startswith('data:'):
#                 base64_file_content = base64_file_content.split(',')[1]
#             try:
#                 file_content = base64.b64decode(base64_file_content)
#             except Exception as e:
#                 return Response({'error': f'Invalid base64 file content for document at index {index}: {str(e)}'}, status=status.HTTP_400_BAD_REQUEST)

#             document_type = document_type.lower()
#             document_text = ""

#             # Process document content based on type
#             if document_type == 'pdf':
#                 pdf_reader = PdfReader(io.BytesIO(file_content))
#                 for page_num in range(len(pdf_reader.pages)):
#                     document_text += pdf_reader.pages[page_num].extract_text() + "\n\n"

#             elif document_type == 'docx':
#                 doc_stream = io.BytesIO(file_content)
#                 doc = Document(doc_stream)
#                 document_text += "\n\n".join([para.text for para in doc.paragraphs]) + "\n\n"

#             elif document_type == 'text':
#                 document_text += file_content.decode('utf-8') + "\n\n"
            

#             elif document_type == 'html':
#                 try:
#                     document_text = self.extract_text_from_html(file_content, document_name)
#                 except Exception as e:
#                     return Response({'error': f'Failed to process HTML for document {document_name}: {str(e)}'}, status=status.HTTP_400_BAD_REQUEST)

#                 document_text += file_content.decode('utf-8') + "\n\n"


#             else:
#                 return Response({'error': f'Unsupported document type: {document_type} for document {document_name}'}, status=status.HTTP_400_BAD_REQUEST)

#             # Save file to S3
#             try:
#                 file_io = io.BytesIO(file_content)
#                 file_io.seek(0)
#                 in_memory_file = InMemoryUploadedFile(
#                     file_io,
#                     None,
#                     document_name,
#                     'application/octet-stream',
#                     len(file_content),
#                     None
#                 )
#                 s3_key = f"documents/{uuid.uuid4()}_{document_name}"
#                 s3_client.upload_fileobj(in_memory_file, settings.AWS_STORAGE_BUCKET_NAME, s3_key)
#                 s3_url = f"https://{settings.AWS_STORAGE_BUCKET_NAME}.s3.{settings.AWS_S3_REGION_NAME}.amazonaws.com/{s3_key}"
#                 s3_urls.append(s3_url)
#             except Exception as e:
#                 return Response({'error': f'Failed to upload document to S3: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

#             # Now, for each document, chunk the text and insert chunks as vectors in Pinecone
#             try:
#                 print("docuemnt text is",document_text)
#                 text_splitter = RecursiveCharacterTextSplitter(separators=separator,chunk_size=chunk_size, chunk_overlap=chunk_overlap)
#                 chunks = text_splitter.split_text(document_text)
#                 # print(chunks)
                
#                 # Embed the chunks into Pinecone as separate vectors
#                 pinecone_initializer.Embeding_Text_list_to_pinecone(
#                     texts=chunks,
#                     index_name=index_name,
#                     namespace=name,
#                     Agent_id=request.user.id,
#                 )
#             except Exception as e:
#                 return Response({'error': f'Failed to insert data into Pinecone for document {document_name}: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

#         # print("checkpoint 6")
#         # If the above succeeds, create the KnowledgeBase
#         try:
#             knowledge_base, created = KnowledgeBase.objects.update_or_create(
#                 name=name,
#                 namespace=name,
#                 created_by=user_admin,
#                 creator=creator,
#                 defaults={
#                     'google_drive_folder_id': folder_id
#                 }
#             )

#             for index, document_data in enumerate(documents):
#                 KnowledgeBaseDocument.objects.update_or_create(
#                     knowledge_base=knowledge_base,
#                     document_name=document_data.get('document_name'),
#                     defaults={
#                         's3_url': s3_urls[index],
#                         'document_type': document_data.get('document_type').lower(),
#                     }
#                 )
#         except Exception as e:
#             return Response({'error': f'Failed to create knowledge base: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
#         # print("checkpoint 7")
#         # Return the successful response with document URLs
#         return Response({
#             'message': 'Knowledge Base created successfully.',
#             'knowledge_base_id': knowledge_base.id,
#             'knowledge_base_name': knowledge_base.name,
#             'documents': s3_urls
#         }, status=status.HTTP_201_CREATED)
    




    # def extract_text_from_html(self,file_content, document_name):
    #     """
    #     Extracts text from HTML content.
    #     """
    #     from bs4 import BeautifulSoup

    #     try:
    #         soup = BeautifulSoup(file_content, 'html.parser')
    #         # Extract text from HTML
    #         text = soup.get_text(separator='\n\n')
    #         return text
    #     except Exception as e:
    #         logger.exception(f"Error extracting text from HTML for {document_name}: {e}")
    #         raise ValueError(f"Error extracting text from HTML: {str(e)}")

class CreateKnowledgeBaseView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        # Retrieve basic data
        name = request.data.get('name')
        print(name)
        admin_id = request.data.get('admin_id')
        folder_id = request.data.get('folder_id')
        chunk_size = request.data.get('chunk_size', 100)
        chunk_overlap = request.data.get('chunk_overlap', 20)
        separator = request.data.get('separator', ["\n\n", "\n", " ", ""])

        print("checkpoint 1")

        if not name or not admin_id:
            return Response({'error': 'Missing required fields: name and admin_id'}, status=status.HTTP_400_BAD_REQUEST)

        # Check permissions and get user admin
        user = request.user
        if user.is_user_admin:
            user_admin = user
            creator = user
        else:
            if user.role != 'editor':
                return Response({'error': 'Only User Admins or Editors can create knowledge bases'}, status=status.HTTP_403_FORBIDDEN)
            try:
                user_admin = User.objects.get(id=admin_id, is_user_admin=True)
            except User.DoesNotExist:
                return Response({'error': 'Admin not found'}, status=status.HTTP_404_NOT_FOUND)

            # Ensure the editor is associated with the admin
            if not UserAdminUserRelationship.objects.filter(user=user, user_admin=user_admin).exists():
                return Response({'error': 'You are not associated with the provided admin.'}, status=status.HTTP_403_FORBIDDEN)

            creator = user
        print("checkpoint 2")
        # Check if the admin already has a Pinecone index, create if needed
        api_settings = APISettings.objects.first()
        if not api_settings:
                raise ImproperlyConfigured("APISettings instance not found.")
            # Initialize Pinecone and clear the existing namespace
        OPENAI_API_KEY = api_settings.openai_api_key or os.getenv("OPENAI_API_KEY")
        PINECONE_API_KEY = api_settings.pinecone_api_key or os.getenv("PINECONE_API_KEY")
        pinecone_initializer = PineconeInitializer(pinecone_api=PINECONE_API_KEY, open_ai_api=OPENAI_API_KEY)
        if not user_admin.pinecone_index:
            index_name = f"index-{user_admin.id}"
            print("usama initialize kr purana index he kr rha hun")
            pinecone_initializer.initialize_pinecone(index_name)
            print("usama puranay index k initialize  se baahr a  rha hun")
            user_admin.pinecone_index = index_name
            user_admin.save()
        else:
            index_name = user_admin.pinecone_index
        # print("checkpoint 3")
        # Check if the namespace already exists
        if KnowledgeBase.objects.filter(created_by=user_admin, namespace=name).exists():
            return Response({'error': 'Namespace already exists for this admin.'}, status=status.HTTP_400_BAD_REQUEST)
        # print("checkpoint 4")
        # Initialize for tracking S3 URLs
        s3_urls = []

        documents = request.data.get('documents', [])
        # print("checkpoint 5")
        for index, document_data in enumerate(documents):
            base64_file_content = document_data.get('file')  # base64 string
            document_type = document_data.get('document_type')
            document_name = document_data.get('document_name')

            if not document_type or not base64_file_content:
                return Response({'error': f'Missing file or document_type for document at index {index}'}, status=status.HTTP_400_BAD_REQUEST)

            # Decode base64 content to bytes
            if base64_file_content.startswith('data:'):
                base64_file_content = base64_file_content.split(',')[1]
            try:
                file_content = base64.b64decode(base64_file_content)
            except Exception as e:
                return Response({'error': f'Invalid base64 file content for document at index {index}: {str(e)}'}, status=status.HTTP_400_BAD_REQUEST)

            document_type = document_type.lower()
            document_text = ""

            # Process document content based on type
            if document_type == 'pdf':
                pdf_reader = PdfReader(io.BytesIO(file_content))
                for page_num in range(len(pdf_reader.pages)):
                    document_text += pdf_reader.pages[page_num].extract_text() + "\n\n"

            elif document_type == 'docx':
                doc_stream = io.BytesIO(file_content)
                doc = Document(doc_stream)
                document_text += "\n\n".join([para.text for para in doc.paragraphs]) + "\n\n"

            elif document_type == 'text':
                document_text += file_content.decode('utf-8') + "\n\n"
            

            elif document_type == 'html':
                try:
                    document_text = self.extract_text_from_html(file_content, document_name)
                except Exception as e:
                    return Response({'error': f'Failed to process HTML for document {document_name}: {str(e)}'}, status=status.HTTP_400_BAD_REQUEST)

                document_text += file_content.decode('utf-8') + "\n\n"


            else:
                return Response({'error': f'Unsupported document type: {document_type} for document {document_name}'}, status=status.HTTP_400_BAD_REQUEST)

            # Save file to S3
            try:
                file_io = io.BytesIO(file_content)
                file_io.seek(0)
                in_memory_file = InMemoryUploadedFile(
                    file_io,
                    None,
                    document_name,
                    'application/octet-stream',
                    len(file_content),
                    None
                )
                s3_key = f"documents/{uuid.uuid4()}_{document_name}.{document_type}"
                s3_client.upload_fileobj(in_memory_file, settings.AWS_STORAGE_BUCKET_NAME, s3_key)
                s3_url = f"https://{settings.AWS_STORAGE_BUCKET_NAME}.s3.{settings.AWS_S3_REGION_NAME}.amazonaws.com/{s3_key}"
                s3_urls.append(s3_url)
                print("s3_urls", s3_urls)
            except Exception as e:
                return Response({'error': f'Failed to upload document to S3: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            # Now, for each document, chunk the text and insert chunks as vectors in Pinecone
            try:
                print("document text is", document_text)
                text_splitter = RecursiveCharacterTextSplitter(
                    separators=separator, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
                chunks = text_splitter.split_text(document_text)
                # For each chunk, prepare individual metadata with specific text content
                s3_url = s3_urls[index] 
                chunk_metadata = [
                    {
                        "document_name": document_name,
                        "s3_url": s3_url,  # Set to the s3 URL of the uploaded document
                        "text": chunk_text.strip()  # Include specific text of this chunk
                    }
                    for chunk_text in chunks
                ]

                # api_settings = APISettings.objects.first()
                # if not api_settings:
                #         raise ImproperlyConfigured("APISettings instance not found.")
                #     # Initialize Pinecone and clear the existing namespace
                # OPENAI_API_KEY = api_settings.openai_api_key or os.getenv("OPENAI_API_KEY")
                # PINECONE_API_KEY = api_settings.pinecone_api_key or os.getenv("PINECONE_API")
                # # print("PINECONE_API_KEY",OPENAI_API_KEY)
                # pinecone_initializer = PineconeInitializer(pinecone_api=PINECONE_API_KEY, open_ai_api=OPENAI_API_KEY)
                # print("PINECONE_API_KEY",PINECONE_API_KEY)
                # Embed the chunks into Pinecone with metadata
                pinecone_initializer.Embeding_Text_list_to_pinecone(
                    texts=chunks,
                    index_name=index_name,
                    namespace=name,
                    Agent_id=request.user.id,
                    metadata_list=chunk_metadata  # Add metadata to each chunk
                )
            except Exception as e:
                return Response({'error': f'Failed to insert data into Pinecone for document {document_name}: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        # print("checkpoint 6")
        # If the above succeeds, create the KnowledgeBase
        try:
            print(name, user_admin, creator)
            if folder_id:
                k_type="drive"
            else:
                k_type="local"
            knowledge_base, created = KnowledgeBase.objects.update_or_create(
                name=name,
                namespace=name,
                created_by=user_admin,
                creator=creator,
                k_type=k_type,
                defaults={
                    'google_drive_folder_id': folder_id
                }
            )
            print(knowledge_base, "reached here")
            for index, document_data in enumerate(documents):
                KnowledgeBaseDocument.objects.update_or_create(
                    knowledge_base=knowledge_base,
                    document_name=document_data.get('document_name'),
                    defaults={
                        's3_url': s3_urls[index],
                        'document_type': document_data.get('document_type').lower(),
                    }
                )
        except Exception as e:
            return Response({'error': f'Failed to create knowledge base: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
        # print("checkpoint 7") THIS 
        # Return the successful response with document URLs
        return Response({
            'message': 'Knowledge Base created successfully.',
            'knowledge_base_id': knowledge_base.id,
            'knowledge_base_name': knowledge_base.name,
            'documents': s3_urls
        }, status=status.HTTP_201_CREATED)
    

    def extract_text_from_html(self,file_content, document_name):
        """
        Extracts text from HTML content.
        """
        from bs4 import BeautifulSoup

        try:
            soup = BeautifulSoup(file_content, 'html.parser')
            # Extract text from HTML
            text = soup.get_text(separator='\n\n')
            return text
        except Exception as e:
            logger.exception(f"Error extracting text from HTML for {document_name}: {e}")
            raise ValueError(f"Error extracting text from HTML: {str(e)}")


class RecreatePineConeIndexView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        # Retrieve request data
        admin_id = request.data.get('admin_id')
        chunk_size = request.data.get('chunk_size', 100)
        chunk_overlap = request.data.get('chunk_overlap', 20)
        separator = request.data.get('separator', ["\n\n", "\n", " ", ""])

        if not admin_id:
            return Response({"error": "admin_id is required."}, status=status.HTTP_400_BAD_REQUEST)

        try:
            # Verify permissions
            user = request.user
            if not user.is_super_admin and (not user.is_user_admin or user.id != admin_id):
                return Response({"error": "Permission denied"}, status=status.HTTP_403_FORBIDDEN)
            
            # Initialize Pinecone for the admin user
            admin_user = User.objects.get(id=admin_id)
            api_settings = APISettings.objects.first()
            if not api_settings:
                raise ImproperlyConfigured("APISettings instance not found.")

            pinecone_initializer = PineconeInitializer(
                pinecone_api=api_settings.pinecone_api_key or os.getenv("PINECONE_API_KEY"),
                open_ai_api=api_settings.openai_api_key or os.getenv("OPENAI_API_KEY")
            )
            
            # Set up or recreate Pinecone index
            index_name = admin_user.pinecone_index or f"index-{admin_user.id}"
            if index_name in pinecone_initializer.show_indexes():
                return Response({"error": "Index already exists."}, status=status.HTTP_400_BAD_REQUEST)
            
            pinecone_initializer.initialize_pinecone(index_name)

            # Recreate KnowledgeBase and re-embed all documents
            knowledge_bases = KnowledgeBase.objects.filter(created_by=admin_user)
            print(knowledge_bases)
            for kb in knowledge_bases:
                documents = KnowledgeBaseDocument.objects.filter(knowledge_base=kb)
                for document in documents:
                    # Fetch and process document content
                    document_text = self.extract_document_text_from_s3(document)
                    
                    # Split the text into chunks with specified chunk size and overlap
                    text_splitter = RecursiveCharacterTextSplitter(
                        separators=separator, chunk_size=chunk_size, chunk_overlap=chunk_overlap
                    )
                    chunks = text_splitter.split_text(document_text)

                    # Prepare metadata for each chunk
                    chunk_metadata = [
                        {
                            "document_name": document.document_name,
                            "s3_url": document.s3_url,
                            "text": chunk_text.strip()
                        }
                        for chunk_text in chunks
                    ]
                    
                    # Embed each chunk with metadata in Pinecone
                    pinecone_initializer.Embeding_Text_list_to_pinecone(
                        texts=chunks,
                        index_name=index_name,
                        namespace=kb.namespace,
                        Agent_id=admin_user.id,
                        metadata_list=chunk_metadata
                    )

            return Response({"success": "Index and knowledge bases successfully recreated."}, status=status.HTTP_200_OK)

        except User.DoesNotExist:
            return Response({"error": "Admin user not found"}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            logger.exception("Error recreating Pinecone index: %s", str(e))
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def extract_document_text_from_s3(self, document):
        """
        Download the document from S3 and extract text based on its type.
        """
        # Initialize document_text to avoid returning None
        document_text = ""
        
        # Download the file content
        file_content = self.download_file_from_s3(document.s3_url)
        document_type = document.document_type.lower()

        # Extract text based on the file type
        if document_type == 'pdf':
            pdf_reader = PdfReader(io.BytesIO(file_content))
            for page in pdf_reader.pages:
                document_text += page.extract_text() + "\n\n"
        elif document_type == 'docx':
            doc_stream = io.BytesIO(file_content)
            doc = Document(doc_stream)
            document_text = "\n\n".join([para.text for para in doc.paragraphs]) + "\n\n"
        elif document_type == 'text':
            document_text = file_content.decode('utf-8') + "\n\n"
        elif document_type == 'html':
            try:
                document_text = self.extract_text_from_html(file_content)
            except Exception as e:
                logger.error(f"Failed to process HTML for document {document.document_name}: {str(e)}")
        else:
            logger.error(f"Unsupported document type: {document_type} for document {document.document_name}")

        return document_text


    def download_file_from_s3(self, s3_url):
        """
        Download file content from S3 using the URL.
        """
        bucket_name = settings.AWS_STORAGE_BUCKET_NAME
        key = s3_url.split(f"https://{bucket_name}.s3.{settings.AWS_S3_REGION_NAME}.amazonaws.com/")[1]

        # Download file content to memory
        with io.BytesIO() as file_io:
            s3_client.download_fileobj(bucket_name, key, file_io)
            file_io.seek(0)
            return file_io.read()

    def extract_text_from_html(self, file_content):
        """
        Extract text from HTML content.
        """
        try:
            soup = BeautifulSoup(file_content, 'html.parser')
            return soup.get_text(separator='\n\n')
        except Exception as e:
            logger.error(f"Error extracting text from HTML: {e}")
            raise ValueError(f"Error extracting text from HTML: {str(e)}")




class DeleteKnowledgeBaseView(APIView):
    permission_classes = [IsAuthenticated]

    def delete(self, request, knowledge_base_id):
        user = request.user  # Can be an Admin or Editor

        try:
            # Retrieve the knowledge base by ID
            knowledge_base = get_object_or_404(KnowledgeBase, id=knowledge_base_id)

            # Check if the user has permission to delete the knowledge base
            if user.is_user_admin and knowledge_base.created_by == user:
                # Admin is deleting their own knowledge base
                pass
            elif user.role == 'editor' and knowledge_base.creator == user:
                # Ensure the editor is associated with the correct admin
                if not UserAdminUserRelationship.objects.filter(user=user, user_admin=knowledge_base.created_by).exists():
                    return Response({'error': 'You do not have permission to delete this knowledge base'}, status=status.HTTP_403_FORBIDDEN)
            else:
                # Neither Admin nor Editor has permission
                return Response({'error': 'You do not have permission to delete this knowledge base'}, status=status.HTTP_403_FORBIDDEN)

            # Initialize Pinecone with the API keys
            api_settings = APISettings.objects.first()
            if not api_settings:
                raise ImproperlyConfigured("APISettings instance not found.")
            OPENAI_API_KEY = api_settings.openai_api_key or os.getenv("OPENAI_API_KEY")
            PINECONE_API_KEY = api_settings.pinecone_api_key or os.getenv("PINECONE_API_KEY")
            pinecone_initializer = PineconeInitializer(pinecone_api=PINECONE_API_KEY, open_ai_api=OPENAI_API_KEY)

            # Get the admin's Pinecone index and namespace
            index_name = knowledge_base.created_by.pinecone_index
            namespace = knowledge_base.namespace

            if index_name:
                # Check if the namespace exists in the Pinecone index
                if not pinecone_initializer.namespace_exists_in_index(index_name=index_name, namespace=namespace):
                    # Namespace does not exist in Pinecone, delete the knowledge base from the database
                    knowledge_base.delete()
                    return Response({'message': 'Knowledge Base deleted from database as the namespace did not exist in Pinecone.'}, status=status.HTTP_200_OK)

                # If namespace exists, delete it from Pinecone
                pinecone_initializer.delete_namespace_from_pinecone(index_name=index_name, namespace=namespace)
            else:
                return Response({'error': 'The admin does not have a Pinecone index.'}, status=status.HTTP_400_BAD_REQUEST)

            # Delete associated documents
            documents = KnowledgeBaseDocument.objects.filter(knowledge_base=knowledge_base)
            for document in documents:
                # Delete the file from S3 or storage
                default_storage.delete(document.s3_url)

            # Delete all document entries related to this knowledge base
            documents.delete()

            # Delete the KnowledgeBase entry from the database
            knowledge_base.delete()

            return Response({'message': 'Knowledge Base and associated namespace deleted successfully.'}, status=status.HTTP_200_OK)

        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

 
class KnowledgeBaseDetailView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, knowledge_base_id):
        user = request.user

        # Retrieve the specific knowledge base from the database
        knowledge_base = get_object_or_404(KnowledgeBase, id=knowledge_base_id)

        # Ensure only user admins or associated editors can access this endpoint
        if not (
                user == knowledge_base.created_by or
                (
                    UserAdminUserRelationship.objects.filter(user=user, user_admin=knowledge_base.created_by).exists() and
                    user.role == 'editor'  # Ensure user is an editor associated with the admin
                )
        ):
            return Response(
                {'error': 'Only associated Editors or User Admins can view this knowledge base'},
                status=status.HTTP_403_FORBIDDEN
            )

        try:
            api_settings = APISettings.objects.first()
            if not api_settings:
                raise ImproperlyConfigured("APISettings instance not found.")

            # Initialize Pinecone
            OPENAI_API_KEY = api_settings.openai_api_key or os.getenv("OPENAI_API_KEY")
            PINECONE_API_KEY = api_settings.pinecone_api_key or os.getenv("PINECONE_API_KEY")
            pinecone_initializer = PineconeInitializer(
                pinecone_api=PINECONE_API_KEY,
                open_ai_api=OPENAI_API_KEY
            )

            # Retrieve Pinecone index information
            index_info = pinecone_initializer.describe_index(knowledge_base.created_by.pinecone_index)
            if index_info is None:
                return Response(
                    {'error': 'Index information could not be retrieved from Pinecone.'},
                    status=status.HTTP_400_BAD_REQUEST
                )

            # Parse Pinecone index information
            dimension = index_info.get('dimension')
            host = index_info.get('host')
            metric = index_info.get('metric')
            name = index_info.get('name')

            # Spec and status sub-dictionaries
            spec_serverless_cloud = index_info['spec']['serverless'].get('cloud')
            spec_serverless_region = index_info['spec']['serverless'].get('region')
            status_ready = index_info['status'].get('ready')
            status_state = index_info['status'].get('state')

            # Get associated documents
            documents = KnowledgeBaseDocument.objects.filter(knowledge_base=knowledge_base)
            document_data = [{
                'document_name': doc.document_name,
                's3_url': doc.s3_url,
                'document_type': doc.document_type,
                'uploaded_at': doc.uploaded_at
            } for doc in documents]

            # Prepare the response data
            response_data = {
                'id': knowledge_base.id,
                'name': knowledge_base.name,
                'namespace': knowledge_base.namespace,
                'k_type': knowledge_base.k_type,
                'created_at': knowledge_base.created_at,
                'created_by': {
                    'id': knowledge_base.created_by.id,
                    'username': knowledge_base.created_by.username,
                    'email': knowledge_base.created_by.email
                },
                'creator': {
                    'id': knowledge_base.creator.id if knowledge_base.creator else None,
                    'username': knowledge_base.creator.username if knowledge_base.creator else None,
                    'email': knowledge_base.creator.email if knowledge_base.creator else None
                },
                'documents': document_data,
                'pinecone_index_info': {
                    'dimension': dimension,
                    'host': host,
                    'metric': metric,
                    'name': name,
                    'spec_serverless_cloud': spec_serverless_cloud,
                    'spec_serverless_region': spec_serverless_region,
                    'status_ready': status_ready,
                    'status_state': status_state,
                }
            }

            return Response(response_data, status=status.HTTP_200_OK)

        except Exception as e:
            # Log the error and return a 500 response if something goes wrong
            logger.error(f"Error retrieving index info from Pinecone: {str(e)}", exc_info=True)
            return Response(
                {'error': f"Error retrieving Pinecone index info: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


class UpdateKnowledgeBaseView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, knowledge_base_id):
        user = request.user
        name = request.data.get('name')
        documents = request.data.get('documents', [])  # List of document objects
        print(documents)
        chunk_size = request.data.get('chunk_size', 100)
        chunk_overlap = request.data.get('chunk_overlap', 20)
        separator = request.data.get('separator', ["\n\n", "\n", " ", ""])

        # Get the knowledge base by ID
        knowledge_base = get_object_or_404(KnowledgeBase, id=knowledge_base_id)

        # Ensure permissions
        if not (
            user == knowledge_base.created_by or
            (
                UserAdminUserRelationship.objects.filter(user=user, user_admin=knowledge_base.created_by).exists() and
                user.role == 'editor'
            )
        ):
            return Response(
                {'error': 'Only associated Editors or User Admins can update this knowledge base'}, 
                status=status.HTTP_403_FORBIDDEN
            )

        try:
            api_settings = APISettings.objects.first()
            if not api_settings:
                raise ImproperlyConfigured("APISettings instance not found.")

            OPENAI_API_KEY = api_settings.openai_api_key or os.getenv("OPENAI_API_KEY")
            PINECONE_API_KEY = api_settings.pinecone_api_key or os.getenv("PINECONE_API_KEY")
            pinecone_initializer = PineconeInitializer(pinecone_api=PINECONE_API_KEY, open_ai_api=OPENAI_API_KEY)

            # Clear existing namespace from Pinecone
            pinecone_initializer.delete_namespace_from_pinecone(
                index_name=knowledge_base.created_by.pinecone_index,
                namespace=knowledge_base.namespace
            )
        except Exception as e:
            return Response({'error': str(e), 'reason': "failed to insert document in index"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        try:
            text_content = ""
            s3_urls = []

            KnowledgeBaseDocument.objects.filter(knowledge_base=knowledge_base).delete()

            # Process new documents
            for index, doc in enumerate(documents):
                document_name = doc.get('document_name')
                document_type = doc.get('document_type', '').lower()
                print(doc)
                # Handle documents with "s3_url" (existing files) or "file" (new files)
                if 's3_url' in doc:
                    s3_url = doc['s3_url']
                    s3_urls.append(s3_url)

                    try:
                        file_content = self.extract_document_text_from_s3(s3_url, document_type, document_name)
                        text_content += file_content
                    except Exception as e:
                        return Response({'error': f'Failed to retrieve or process document from S3 URL: {str(e)}'}, status=status.HTTP_400_BAD_REQUEST)

                elif 'file' in doc:
                    base64_file_content = doc['file']
                    
                    # Strip out the `data:<mime-type>;base64,` part if present
                    if base64_file_content.startswith('data:'):
                        base64_file_content = base64_file_content.split(',')[1]

                    # Decode base64 content
                    try:
                        file_content = base64.b64decode(base64_file_content)
                    except Exception as e:
                        return Response(
                            {'error': f'Invalid base64 file content for document at index {index}: {str(e)}'}, 
                            status=status.HTTP_400_BAD_REQUEST
                        )

                    # Process document based on type
                    if document_type == 'pdf':
                        pdf_reader = PdfReader(io.BytesIO(file_content))
                        for page_num in range(len(pdf_reader.pages)):
                            text_content += pdf_reader.pages[page_num].extract_text()

                    elif document_type == 'docx':
                        doc_stream = io.BytesIO(file_content)
                        doc = Document(doc_stream)
                        text_content += "\n".join([para.text for para in doc.paragraphs])

                    elif document_type == 'text':
                        text_content += file_content.decode('utf-8')

                    elif document_type == 'html':
                        try:
                            text_content = self.extract_text_from_html(file_content)
                        except Exception as e:
                            return Response(
                                {'error': f'Failed to process HTML for document {document_name}: {str(e)}'},
                                status=status.HTTP_400_BAD_REQUEST
                            )

                    # Save to S3
                    try:
                        file_io = io.BytesIO(file_content)
                        file_io.seek(0)
                        s3_key = f"documents/{uuid.uuid4()}_{document_name}.{document_type}"
                        s3_client.upload_fileobj(file_io, settings.AWS_STORAGE_BUCKET_NAME, s3_key)
                        s3_url = f"https://{settings.AWS_STORAGE_BUCKET_NAME}.s3.{settings.AWS_S3_REGION_NAME}.amazonaws.com/{s3_key}"
                        s3_urls.append(s3_url)
                    except Exception as e:
                        return Response({'error': f'Failed to upload document: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

                else:
                    return Response({'error': 'Document must contain either "s3_url" or "file" field'}, status=status.HTTP_400_BAD_REQUEST)

                # Create a record in KnowledgeBaseDocument model
                KnowledgeBaseDocument.objects.create(
                    knowledge_base=knowledge_base,
                    document_name=document_name,
                    s3_url=s3_url,
                    document_type=document_type
                )

            # Embed the chunks into Pinecone
            text_splitter = RecursiveCharacterTextSplitter(separators=separator, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
            chunks = text_splitter.split_text(text_content)
            chunk_metadata = [
                {
                    "document_name": document_name,
                    "s3_url": s3_url,
                    "text": chunk_text.strip()
                }
                for chunk_text in chunks
            ]
            pinecone_initializer.Embeding_Text_list_to_pinecone(
                texts=chunks,
                index_name=knowledge_base.created_by.pinecone_index,
                namespace=knowledge_base.namespace,
                Agent_id=knowledge_base.created_by.id,
                metadata_list=chunk_metadata
            )

            return Response({
                'message': 'Knowledge Base updated successfully.',
                'knowledge_base_id': knowledge_base.id,
                'knowledge_base_name': knowledge_base.name,
                'documents': s3_urls
            }, status=status.HTTP_200_OK)

        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def extract_document_text_from_s3(self, s3_url, document_type, document_name):
        """
        Download the document from S3 and extract text based on its type.
        """
        document_text = ""
        file_content = self.download_file_from_s3(s3_url)
        # print("document.s3_url")
        # document_type = document.document_type.lower()

        # Extract text based on the file type
        if document_type == 'pdf':
            pdf_reader = PdfReader(io.BytesIO(file_content))
            for page in pdf_reader.pages:
                document_text += page.extract_text() + "\n\n"
        elif document_type == 'docx':
            doc_stream = io.BytesIO(file_content)
            doc = Document(doc_stream)
            document_text = "\n\n".join([para.text for para in doc.paragraphs]) + "\n\n"
        elif document_type == 'text':
            document_text = file_content.decode('utf-8') + "\n\n"
        elif document_type == 'html':
            try:
                document_text = self.extract_text_from_html(file_content)
            except Exception as e:
                logger.error(f"Failed to process HTML for document {document_name}: {str(e)}")
        else:
            logger.error(f"Unsupported document type: {document_type} for document {document_name}")

        return document_text


    def download_file_from_s3(self, s3_url):
        """
        Download file content from S3 using the URL.
        """
        # Parse the S3 URL to extract the bucket and key
        parsed_url = urlparse(s3_url)
        bucket_name = settings.AWS_STORAGE_BUCKET_NAME
        key = parsed_url.path.lstrip('/')  # Strip the leading '/' from the path

        # Download file content to memory
        with io.BytesIO() as file_io:
            s3_client.download_fileobj(bucket_name, key, file_io)
            file_io.seek(0)
            return file_io.read()

    def extract_text_from_html(self, file_content):
        """
        Extracts text from HTML content.
        """
        try:
            soup = BeautifulSoup(file_content, 'html.parser')
            return soup.get_text(separator='\n\n')
        except Exception as e:
            logger.error(f"Error extracting text from HTML: {e}")
            raise ValueError(f"Error extracting text from HTML: {str(e)}")




class CheckUserPineconeIndexView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        user = request.user

        # Retrieve the Pinecone index name from the user's profile
        pinecone_index_name = user.pinecone_index
        # print(pinecone_index_name)
        # pinecone_index_name = pinecone_index_name.strip()
        # # print(pinecone_index_name)
        if not pinecone_index_name:
            return Response({'error': 'No Pinecone index assigned to this user.'}, status=status.HTTP_404_NOT_FOUND)

        # Check if the Pinecone index exists
        api_settings = APISettings.objects.first()
        if not api_settings:
            return Response({'error': 'API settings not configured.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        OPENAI_API_KEY = api_settings.openai_api_key or os.getenv("OPENAI_API_KEY")
        PINECONE_API_KEY = api_settings.pinecone_api_key or os.getenv("PINECONE_API_KEY")
        pinecone_initializer = PineconeInitializer(pinecone_api=PINECONE_API_KEY, open_ai_api=OPENAI_API_KEY)

        # List existing indexes in Pinecone to verify if the user's index exists
        try:
            existing_indexes = pinecone_initializer.show_indexes()  # Get a list of all existing Pinecone indexes
            index_names = [index['name'] for index in existing_indexes]
            print("existing indexes",index_names)

            # Check if the user's index is in the list of existing indexes
            if pinecone_index_name in index_names:
                return Response({
                    'message': f'Pinecone index "{pinecone_index_name}" exists for the user.'
                }, status=status.HTTP_200_OK)
            else:
                return Response({
                    'error':f'Pinecone index does not exist.',
                    'some':f'{pinecone_index_name} with type is {type(pinecone_index_name)}',
                },status=status.HTTP_404_NOT_FOUND)
        
        except Exception as e:
            return Response({'error': f'Error checking Pinecone index: {str(e)}','some':f'{pinecone_index_name} with type is {type(pinecone_index_name)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)



class CreateEmptyPineConeIndexView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        # Retrieve the admin_id from the request data
        admin_id = request.data.get('admin_id')

        if not admin_id:
            return Response({"error": "admin_id is required."}, status=status.HTTP_400_BAD_REQUEST)

        try:
            # Verify permissions
            user = request.user
            if not user.is_super_admin and (not user.is_user_admin or user.id != admin_id):
                return Response({"error": "Permission denied"}, status=status.HTTP_403_FORBIDDEN)
            
            # Initialize Pinecone for the admin user
            admin_user = User.objects.get(id=admin_id)
            api_settings = APISettings.objects.first()
            if not api_settings:
                raise ImproperlyConfigured("APISettings instance not found.")

            pinecone_initializer = PineconeInitializer(
                pinecone_api=api_settings.pinecone_api_key or os.getenv("PINECONE_API_KEY"),
                open_ai_api=api_settings.openai_api_key or os.getenv("OPENAI_API_KEY")
            )
            
            # Set up Pinecone index with the existing name, or create one if it doesn't exist
            index_name = admin_user.pinecone_index or f"index-{admin_user.id}"
            
            # Delete the existing index if it exists
            existing_indexes = pinecone_initializer.show_indexes()
            index_names = [index['name'] for index in existing_indexes]
            if index_name in index_names:
                pinecone_initializer.delete_index_pinecone(index_name)
            
            # Recreate an empty index in Pinecone
            pinecone_initializer.initialize_pinecone(index_name)

            # Delete all knowledge bases and documents associated with this admin user
            KnowledgeBase.objects.filter(created_by=admin_user).delete()
            KnowledgeBaseDocument.objects.filter(knowledge_base__created_by=admin_user).delete()

            return Response({"success": f"New empty Pinecone index '{index_name}' created with no knowledge bases or documents."}, status=status.HTTP_200_OK)

        except User.DoesNotExist:
            return Response({"error": "Admin user not found"}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            logger.exception("Error creating empty Pinecone index: %s", str(e))
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class ListKnowledgeBasesView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        user_admin_id = request.query_params.get('user_admin_id')
        user = request.user  # Can be an Admin or an Editor

        # If user_admin_id is provided, fetch the knowledge bases for that admin
        if user_admin_id:
            try:
                user_admin = User.objects.get(id=user_admin_id, is_user_admin=True)
            except User.DoesNotExist:
                return Response({'error': 'User Admin not found'}, status=status.HTTP_404_NOT_FOUND)

            # Get all knowledge bases created by the specified user admin
            knowledge_bases = KnowledgeBase.objects.filter(created_by=user_admin)

        else:
            # If no user_admin_id is provided, check the current user's role
            if user.is_user_admin:
                # Admin: Fetch all knowledge bases created by the current admin
                knowledge_bases = KnowledgeBase.objects.filter(created_by=user)
            elif user.role == 'editor':
                # Editor: Fetch all knowledge bases associated with this editor's admin(s)
                admin_ids = UserAdminUserRelationship.objects.filter(user=user).values_list('user_admin', flat=True)
                knowledge_bases = KnowledgeBase.objects.filter(created_by__in=admin_ids)
            else:
                return Response({'error': 'You do not have permission to view knowledge bases'}, status=status.HTTP_403_FORBIDDEN)


        api_settings = APISettings.objects.first()
        if not api_settings:
            raise ImproperlyConfigured("APISettings instance not found.")
        # Initialize Pinecone and clear the existing namespace
        OPENAI_API_KEY = api_settings.openai_api_key or os.getenv("OPENAI_API_KEY")
        PINECONE_API_KEY = api_settings.pinecone_api_key or os.getenv("PINECONE_API_KEY")
        pinecone_initializer = PineconeInitializer(pinecone_api=PINECONE_API_KEY, open_ai_api=OPENAI_API_KEY)

    # Serialize the knowledge bases
        knowledge_bases_data = []
        for kb in knowledge_bases:
            try:
                # # Check if the namespace exists in Pinecone
                # index_name = kb.created_by.pinecone_index
                # namespace_exists = pinecone_initializer.namespace_exists(index_name=index_name, namespace=kb.namespace)

                # # If the namespace does not exist, delete the knowledge base and associated documents
                # if not namespace_exists:
                #     KnowledgeBaseDocument.objects.filter(knowledge_base=kb).delete()  # Delete associated documents
                #     kb.delete()  # Delete knowledge base entry
                #     continue  # Skip this knowledge base in the response

                # Get associated documents for this knowledge base
                documents = KnowledgeBaseDocument.objects.filter(knowledge_base=kb)
                document_data = [{
                    'document_name': doc.document_name,
                    's3_url': doc.s3_url,
                    'document_type': doc.document_type,
                    'uploaded_at': doc.uploaded_at
                } for doc in documents]

                knowledge_base_data = {
                    'id': kb.id,
                    'name': kb.name,
                    'namespace': kb.namespace,
                    'created_at': kb.created_at,
                    'created_by': {
                        'id': kb.created_by.id,
                        'username': kb.created_by.username,
                        'email': kb.created_by.email
                    },
                    'creator': {
                        'id': kb.creator.id if kb.creator else None,
                        'username': kb.creator.username if kb.creator else None,
                        'email': kb.creator.email if kb.creator else None
                    },
                    'documents': document_data
                }
                knowledge_bases_data.append(knowledge_base_data)

            except Exception as e:
                # Log or handle exceptions related to Pinecone checks
                logger.error(f"Error checking namespace in Pinecone for knowledge base {kb.id}: {str(e)}", exc_info=True)

        return Response({'knowledge_bases': knowledge_bases_data}, status=status.HTTP_200_OK)




# # Endpoint for Uploading Files to AWS S3 Bucket
# class UploadFileToS3View(APIView):
#     permission_classes = [IsAuthenticated]

#     def post(self, request):
#         # Retrieve data from request
#         base64_file_content = request.data.get('file')  # base64 string
#         document_name = request.data.get('document_name')

#         # Check if required fields are present
#         if not base64_file_content or not document_name:
#             return Response({'error': 'Missing required fields: file and document_name'}, status=status.HTTP_400_BAD_REQUEST)

#         # Decode base64 content to bytes
#         if base64_file_content.startswith('data:'):
#             base64_file_content = base64_file_content.split(',')[1]
        
#         try:
#             file_content = base64.b64decode(base64_file_content)
#         except Exception as e:
#             return Response({'error': f'Invalid base64 file content: {str(e)}'}, status=status.HTTP_400_BAD_REQUEST)

#         # Save file to S3
#         try:
#             file_io = io.BytesIO(file_content)
#             file_io.seek(0)
#             in_memory_file = InMemoryUploadedFile(
#                 file_io,
#                 None,
#                 document_name,
#                 'application/octet-stream',
#                 len(file_content),
#                 None
#             )
#             s3_key = f"documents/{uuid.uuid4()}_{document_name}"
#             s3_client.upload_fileobj(in_memory_file, settings.AWS_STORAGE_BUCKET_NAME, s3_key)
#             s3_url = f"https://{settings.AWS_STORAGE_BUCKET_NAME}.s3.{settings.AWS_S3_REGION_NAME}.amazonaws.com/{s3_key}"

#             return Response({'message': 'File uploaded successfully', 'file_url': s3_url}, status=status.HTTP_201_CREATED)

#         except Exception as e:
#             return Response({'error': f'Failed to upload file to S3: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)



# Endpoint for Uploading Files to AWS S3 Bucket
class UploadFileToS3View(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        # Retrieve data from request
        base64_file_content = request.data.get('file')  # base64 string
        document_name = request.data.get('document_name')

        # Check if required fields are present
        if not base64_file_content or not document_name:
            return Response({'error': 'Missing required fields: file and document_name'}, status=status.HTTP_400_BAD_REQUEST)

        # Decode base64 content to bytes
        if base64_file_content.startswith('data:'):
            base64_file_content = base64_file_content.split(',')[1]
        
        try:
            file_content = base64.b64decode(base64_file_content)
        except Exception as e:
            return Response({'error': f'Invalid base64 file content: {str(e)}'}, status=status.HTTP_400_BAD_REQUEST)

        # Save file to S3
        try:
            file_io = io.BytesIO(file_content)
            file_io.seek(0)
            in_memory_file = InMemoryUploadedFile(
                file_io,
                None,
                document_name,
                'application/octet-stream',
                len(file_content),
                None
            )
            s3_key = f"documents/{uuid.uuid4()}_{document_name}"
            s3_client.upload_fileobj(in_memory_file, settings.AWS_STORAGE_BUCKET_NAME, s3_key)
            s3_url = f"https://{settings.AWS_STORAGE_BUCKET_NAME}.s3.{settings.AWS_S3_REGION_NAME}.amazonaws.com/{s3_key}"

            return Response({'message': 'File uploaded successfully', 's3_url': s3_url}, status=status.HTTP_201_CREATED)

        except Exception as e:
            return Response({'error': f'Failed to upload file to S3: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)





# google-drive k endpoints

class GoogleDriveInitView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        # Initialize the Google OAuth flow
        flow = Flow.from_client_secrets_file(
            settings.GOOGLE_DRIVE_CREDENTIALS,
            scopes=['https://www.googleapis.com/auth/drive.readonly'],
            redirect_uri=settings.GOOGLE_DRIVE_REDIRECT_URI
        )
        
        # Generate a unique state to store user info
        state = urlencode({'user_id': str(request.user.id), 'session_id': str(uuid.uuid4())})
        
        # Generate the auth URL with the state parameter
        auth_url, _ = flow.authorization_url(prompt='consent', state=state)
        
        # Return the auth URL for the user to follow
        return Response({'auth_url': auth_url}, status=status.HTTP_200_OK)




# class GoogleDriveCallbackView(APIView):
#     permission_classes = [AllowAny]  # Allow access to Google callback without authentication

#     def get(self, request):
#         # Get the authorization code and state from the query parameters
#         code = request.GET.get('code')
#         state = request.GET.get('state')
        
#         if not code or not state:
#             return Response({'error': 'No code or state provided.'}, status=status.HTTP_400_BAD_REQUEST)
        
#         # Extract user_id from the state parameter
#         try:
#             state_data = dict(pair.split('=') for pair in state.split('&'))
#             user_id = state_data.get('user_id')
#         except:
#             return Response({'error': 'Invalid state parameter.'}, status=status.HTTP_400_BAD_REQUEST)
        
#         # Find the user based on the user_id
#         User = get_user_model()
#         try:
#             user = User.objects.get(id=user_id)
#         except User.DoesNotExist:
#             return Response({'error': 'User not found.'}, status=status.HTTP_404_NOT_FOUND)

#         # Initialize the Google OAuth flow
#         flow = Flow.from_client_secrets_file(
#             settings.GOOGLE_DRIVE_CREDENTIALS,
#             scopes=['https://www.googleapis.com/auth/drive.readonly'],
#             redirect_uri=settings.GOOGLE_DRIVE_REDIRECT_URI
#         )
        
#         try:
#             # Fetch tokens using the provided authorization code
#             flow.fetch_token(code=str(code))
#             credentials = flow.credentials

#             # Store the credentials in the database
#             creds_data = {
#                 'token': credentials.token,
#                 'refresh_token': credentials.refresh_token,
#                 'token_uri': credentials.token_uri,
#                 'client_id': credentials.client_id,
#                 'client_secret': credentials.client_secret,
#                 'scopes': credentials.scopes
#             }

#             # Update or create Google Drive account information for the user
#             GoogleDriveAccount.objects.update_or_create(
#                 user=user,
#                 defaults={'credentials': json.dumps(creds_data)}
#             )

#             # Return success response or redirect to your frontend
            # return Response({
            #     'message': 'Google Drive successfully authenticated and tokens saved.',
            #     'credentials': creds_data
            # }, status=status.HTTP_200_OK)

#         except Exception as e:
#             # Catch any other errors that might occur
#             return Response({'error': 'An error occurred: ' + str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class GoogleDriveCallbackView(APIView):
    permission_classes = [AllowAny]

    def get(self, request):
        # Get the authorization code and state from the query parameters
        code = request.GET.get('code')
        state = request.GET.get('state')

        if not code or not state:
            return Response({'error': 'No code or state provided.'}, status=status.HTTP_400_BAD_REQUEST)

        # Extract user_id from the state parameter
        try:
            state_data = dict(pair.split('=') for pair in state.split('&'))
            user_id = state_data.get('user_id')
            # print("User_id=", user_id)
        except:
            return Response({'error': 'Invalid state parameter.'}, status=status.HTTP_400_BAD_REQUEST)

        # Find the user based on the user_id
        User = get_user_model()
        try:
            user = User.objects.get(id=user_id)
            # print("User:", user)
        except User.DoesNotExist:
            return Response({'error': 'User not found.'}, status=status.HTTP_404_NOT_FOUND)

        # Initialize the Google OAuth flow
        flow = Flow.from_client_secrets_file(
            settings.GOOGLE_DRIVE_CREDENTIALS,
            scopes=['https://www.googleapis.com/auth/drive.readonly'],
            redirect_uri=settings.GOOGLE_DRIVE_REDIRECT_URI
        )

        try:
            # Fetch tokens using the provided authorization code
            flow.fetch_token(code=str(code))
            credentials = flow.credentials

            # Store the credentials in the database
            creds_data = {
                'token': credentials.token,
                'refresh_token': credentials.refresh_token,
                'token_uri': credentials.token_uri,
                'client_id': credentials.client_id,
                'client_secret': credentials.client_secret,
                'scopes': credentials.scopes
            }
            # print("Cresds:" ,creds_data)
            GoogleDriveAccount.objects.update_or_create(
                user=user,
                defaults={'credentials': json.dumps(creds_data)}
            )
            # print("grive creds saved successfully...")

            # Generate JWT token for the user
            jwt_payload = {
                'user_id': user.id
            }
            token = jwt.encode(jwt_payload, settings.SECRET_KEY, algorithm='HS256')

            # Redirect to the frontend with JWT and Google access token in query parameters
            frontend_url = 'https://ai-rag-client-admin-git-staging-octalooptechnologies-projects.vercel.app/knowledgebase/own'
            # frontend_url = 'http://localhost:5173/knowledgebase/own'
            params = {
                'token': token,
                'access_token': credentials.token,  # Google access token
                'refresh_token': credentials.refresh_token  # Google refresh token (optional)
            }
            redirect_url = f"{frontend_url}?{urlencode(params)}"
            
            return HttpResponseRedirect(redirect_url)
            # return Response({
            #     'message': 'Google Drive successfully authenticated and tokens saved.',
            #     'credentials': creds_data
            # }, status=status.HTTP_200_OK)

        except Exception as e:
            return Response({'error': 'An error occurred: ' + str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

# class GoogleDriveCallbackView(APIView):
#     permission_classes = [IsAuthenticated]

#     def post(self, request):
#         """
#         Endpoint to save Google Drive credentials for the authenticated user.
#         Expects the 'code' (authorization code) and 'redirect_uri' in the request payload.
#         """
#         # Extract authorization code and redirect URI from request
#         code = request.data.get('code')
#         redirect_uri = request.data.get('redirect_uri')

#         if not code or not redirect_uri:
#             return Response({'error': 'Missing code or redirect_uri in request payload.'}, status=status.HTTP_400_BAD_REQUEST)

#         try:
#             # Exchange authorization code for tokens
#             token_url = "https://oauth2.googleapis.com/token"
#             data = {
#                 "code": code,
#                 "client_id": settings.GOOGLE_CLIENT_ID,
#                 "client_secret": settings.GOOGLE_CLIENT_SECRET,
#                 "redirect_uri": redirect_uri,
#                 "grant_type": "authorization_code",
#             }

#             token_response = requests.post(token_url, data=data)

#             # Check if the request to Google was successful
#             if token_response.status_code != 200:
#                 # Log the full response for debugging
#                 logging.error(f"Google OAuth token exchange failed: {token_response}")

#                 # Return the full error message to help debug the problem
#                 return Response({
#                     'error': 'Failed to exchange authorization code with Google.',
#                     'details': token_response.json(),  # Include Google’s error response,
#                      "code": code,
#                     "uri": redirect_uri
#                 }, status=status.HTTP_400_BAD_REQUEST)

#             # Process the token data if successful
#             token_data = token_response.json()
#             access_token = token_data.get('access_token')
#             refresh_token = token_data.get('refresh_token')
#             token_uri = token_data.get('token_uri', 'https://oauth2.googleapis.com/token')
#             client_id = settings.GOOGLE_CLIENT_ID
#             client_secret = settings.GOOGLE_CLIENT_SECRET
#             scopes = token_data.get('scope', '')

#             # Check if both tokens are retrieved
#             if not access_token or not refresh_token:
#                 return Response({'error': 'Failed to retrieve access and refresh tokens.'}, status=status.HTTP_400_BAD_REQUEST)

#             # Store credentials in the database associated with the user
#             credentials = {
#                 "token": access_token,
#                 "refresh_token": refresh_token,
#                 "token_uri": token_uri,
#                 "client_id": client_id,
#                 "client_secret": client_secret,
#                 "scopes": scopes.split(),
#             }

#             GoogleDriveAccount.objects.update_or_create(
#                 user=request.user,
#                 defaults={'credentials': json.dumps(credentials)}
#             )

#             return Response({'message': 'Google Drive credentials saved successfully.'}, status=status.HTTP_200_OK)

#         except Exception as e:
#             # Handle any errors during the process
#             logging.error(f"An error occurred while saving credentials: {str(e)}")
#             return Response({'error': f'An error occurred while saving credentials: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class GoogleDriveFolderFilesView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        folder_id = request.data.get('folder_id')
        name = request.data.get('name')
        admin_id = request.data.get('admin_id')
        chunk_size = request.data.get('chunk_size', 100)
        chunk_overlap = request.data.get('chunk_overlap', 20)
        separator = request.data.get('separator', ["\n\n", "\n", " ", ""])

        user_admin = request.user
                # Check if the namespace already exists
        if KnowledgeBase.objects.filter(created_by=user_admin, namespace=name).exists():
            return Response({'error': 'Namespace already exists for this admin.'}, status=status.HTTP_400_BAD_REQUEST)

        if not name or not admin_id:
            return Response({'error': 'Missing required fields: name and admin_id'}, status=status.HTTP_400_BAD_REQUEST)

        if not folder_id:
            return Response({'error': 'Missing required field: folder_id'}, status=status.HTTP_400_BAD_REQUEST)

        # Retrieve the credentials for the authenticated user
       # Retrieve the credentials for the authenticated user
        try:
            print(Fore.YELLOW , "in try ")
            google_drive_account = GoogleDriveAccount.objects.get(user=request.user)
            credentials_info = json.loads(google_drive_account.credentials)

            # Check if access token is present and valid
            if 'token' in credentials_info:
                access_token = credentials_info['token']
                # access_token = "ya29.a0AeDClZBoG3epsk3JTygtjN8_8SD5aOaZQrlYrPm1VR--jyUq1g650ZinqZXZSJLgsmwAtpcz-C_zXDZzXnqP66rfPkIaouQtm0DdRtlup8HEpDUxk_imR7vDkxbu8OqLdILkFN95WkJl8vOUOxrBOTSQX71CBJxrcJjPrtYkaCgYKAT8SARESFQHGX2MivN_kfFRqatSjXbCYNh5g_w0175"
            else:
                return Response({'error': 'Access token missing in credentials'}, status=status.HTTP_400_BAD_REQUEST)

            credentials = Credentials(**credentials_info)
            # Create credentials from the access token
            creds = Credentials(token=access_token)

            # If credentials have expired and we have a refresh token, refresh the access token
            if credentials and credentials.expired and credentials.refresh_token:
                credentials.refresh(Request()) 
                                # Save the refreshed credentials back to the database
                updated_credentials = {
                    'token': credentials.token,
                    'refresh_token': credentials.refresh_token,
                    'token_uri': credentials.token_uri,
                    'client_id': credentials.client_id,
                    'client_secret': credentials.client_secret,
                    'scopes': credentials.scopes,
                }
                google_drive_account.credentials = json.dumps(updated_credentials)
                google_drive_account.save()

        except GoogleDriveAccount.DoesNotExist:
            return Response({'error': 'Google Drive credentials not found for user'}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            return Response({'error': f'Failed to load credentials: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
        # Use the credentials to access Google Drive API and list files in the folder
        try:
            print(Fore.YELLOW , "in try 1")

            service = build('drive', 'v3', credentials=creds)
            print(Fore.YELLOW , "in try service  :  ", service)
            query = f"'{folder_id}' in parents and trashed = false"
            print(Fore.YELLOW , "in try query  :  " , query)
            results = service.files().list(q=query, fields="files(id, name, mimeType)").execute()  # <<<
            print(Fore.YELLOW , "results :  ", results)
            print(results)
            files = results.get('files', [])

            if not files:
                return Response({'message': 'No files found in the folder'}, status=status.HTTP_200_OK)

            # Prepare documents for the knowledge base
            documents = []
            for file in files:
                print(Fore.YELLOW , "now file is  " , file)

                file_id = file.get('id')
                file_name = file.get('name')
                mime_type = file.get('mimeType')
                # print("filename", file_name)
                # print(mime_type)

                # Get document type from MIME type
                document_type = self.get_document_type(mime_type)

                # Skip files that are not supported
                if document_type == 'unknown':
                    continue

                # Download the file content
                file_content = service.files().get_media(fileId=file_id).execute()
                # print()

                # Convert file content to base64 for further processing
                base64_file_content = base64.b64encode(file_content).decode('utf-8')
                document = {
                    'file': base64_file_content,
                    'document_name': file_name,
                    'document_type': document_type
                }
                documents.append(document)

            # If no supported files are found, return a message
            if not documents:
                return Response({'message': 'No supported file types found in the folder (pdf, docx, text, html).'}, status=status.HTTP_200_OK)

            # Forward the documents to the knowledge base creation API
            knowledge_base_data = {
                'name': name,
                'admin_id': admin_id,
                'folder_id': folder_id,
                'chunk_size': chunk_size,
                'chunk_overlap': chunk_overlap,
                'separator': separator,
                'documents': documents
            }

            # Forward to CreateKnowledgeBaseView
            print(Fore.YELLOW , "going to farward  ")

            knowledge_base_response = self.forward_to_create_knowledge_base(knowledge_base_data)
            return knowledge_base_response
            # return Response(f"Success: {knowledge_base_data}", status=status.HTTP_200_OK)

        except Exception as e:
            return Response({'error': f'Failed to read files from Google Drive: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def get_document_type(self, mime_type):
        """
        Maps Google Drive MIME types to document types.
        Only returns supported types: pdf, docx, text, html.
        """
        if 'pdf' in mime_type:
            return 'pdf'
        elif 'vnd.openxmlformats-officedocument.wordprocessingml.document' in mime_type:
            return 'docx'
        elif 'plain' in mime_type:
            return 'text'
        elif 'html' in mime_type:
            return 'html'
        else:
            return 'unknown'

    def forward_to_create_knowledge_base(self, knowledge_base_data):
        """
        Forward the base64-encoded documents to the knowledge base creation API.
        """
        create_kb_url = 'https://ragai.octalooptechnologies.com/api/useradmin/knowledge-base/create/' 
        # create_kb_url = 'http://127.0.0.1:8000/api/useradmin/knowledge-base/create/' # Assuming this is the endpoint for creating knowledge base
        # Correct the Authorization header by removing the 'b' prefix and decoding the token if necessary
        headers = {
            'Authorization': f'Bearer {self.request.auth.token.decode("utf-8")}' if isinstance(self.request.auth.token, bytes) else f'Bearer {self.request.auth.token}',
            'Content-Type': 'application/json'
        }
        print("headers",headers)
        print("knowledge_base:", knowledge_base_data)
        try:
            response = requests.post(create_kb_url, json=knowledge_base_data, headers=headers)
            response.raise_for_status()  # Raise an exception for 4XX/5XX responses
            return Response(response.json(), status=response.status_code)
        except requests.exceptions.RequestException as e:
            return Response({'error': f'Failed to create knowledge base: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


# class GoogleDriveFolderFilesView(APIView):
#     # We can now include `permission_classes` if desired since we’re in a sync context
#     # permission_classes = [IsAuthenticaea  ted]

#     def post(self, request):
#         print(Fore.GREEN, "Start working")

#         # Handle authentication manually if needed
#         auth = JWTAuthentication()
#         user, _ = auth.authenticate(request)
#         if user is None:
#             return Response({'error': 'Authentication required'}, status=status.HTTP_401_UNAUTHORIZED)

#         # Process request data
#         folder_id = request.data.get('folder_id')
#         name = request.data.get('name')
#         admin_id = request.data.get('admin_id')
#         chunk_size = request.data.get('chunk_size', 100)
#         chunk_overlap = request.data.get('chunk_overlap', 20)
#         separator = request.data.get('separator', ["\n\n", "\n", " ", ""])
#         print(Fore.GREEN, "After payload")

#         if not name or not admin_id:
#             return Response({'error': 'Missing required fields: name and admin_id'}, status=status.HTTP_400_BAD_REQUEST)

#         if not folder_id:
#             return Response({'error': 'Missing required field: folder_id'}, status=status.HTTP_400_BAD_REQUEST)

#         # Assume an access token is directly available for demonstration
#         access_token = "ya29.a0AeDClZDiinGCT2oNHMDezhmfYThLCxcPWQunvugU8Cig86paRJWhjYQ6EMWbTMFhtiuLY-tvKRS6MRFoZbrpt_ZbGWYV7iwvJkaXe1VZPRIifXUeth2xdZjjTszDExY42_0dGmM2pJD-Oj5J_OdtkgJXLDOQCedqTAP9HVGQaCgYKAQgSARESFQHGX2MiVgM0clyra4cgtO0FIYiMeA0175"
#         creds = Credentials(token=access_token)

#         try:
#             print(Fore.GREEN, "After credential creation in try block")

#             service = build('drive', 'v3', credentials=creds)
#             query = f"'{folder_id}' in parents and trashed = false"
#             results = async_to_sync(service.files().list(q=query, fields="files(id, name, mimeType)").execute)()
#             files = results.get('files', [])

#             if not files:
#                 return Response({'message': 'No files found in the folder'}, status=status.HTTP_200_OK)

#             documents = []
#             for file in files:
#                 print(Fore.GREEN, "Processing file:", file)

#                 file_id = file.get('id')
#                 file_name = file.get('name')
#                 mime_type = file.get('mimeType')

#                 document_type = async_to_sync(self.get_document_type)(mime_type)

#                 if document_type == 'unknown':
#                     continue

#                 file_content = async_to_sync(service.files().get_media(fileId=file_id).execute)()
#                 base64_file_content = base64.b64encode(file_content).decode('utf-8')
#                 documents.append({
#                     'file': base64_file_content,
#                     'document_name': file_name,
#                     'document_type': document_type
#                 })

#             if not documents:
#                 return Response({'message': 'No supported file types found in the folder (pdf, docx, text, html).'}, status=status.HTTP_200_OK)

#             knowledge_base_data = {
#                 'name': name,
#                 'admin_id': admin_id,
#                 'folder_id': folder_id,
#                 'chunk_size': chunk_size,
#                 'chunk_overlap': chunk_overlap,
#                 'separator': separator,
#                 'documents': documents
#             }

#             print(Fore.GREEN, "Forwarding knowledge base data")
#             knowledge_base_response = async_to_sync(self.forward_to_create_knowledge_base)(knowledge_base_data)
#             return knowledge_base_response

#         except Exception as e:
#             return Response({'error': f'Failed to read files from Google Drive: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

#     def get_document_type(self, mime_type):
#         if 'pdf' in mime_type:
#             return 'pdf'
#         elif 'vnd.openxmlformats-officedocument.wordprocessingml.document' in mime_type:
#             return 'docx'
#         elif 'plain' in mime_type:
#             return 'text'
#         elif 'html' in mime_type:
#             return 'html'
#         else:
#             return 'unknown'

#     def forward_to_create_knowledge_base(self, knowledge_base_data):
#         create_kb_url = 'http://127.0.0.1:8000/api/useradmin/knowledge-base/create/'
#         headers = {
#             'Authorization': f'Bearer {self.request.auth.token}',
#             'Content-Type': 'application/json'
#         }

#         print("Headers:", headers)
#         print("Knowledge Base Data:", knowledge_base_data)

#         try:
#             # Use aiohttp for async HTTP request wrapped with async_to_sync
#             async def async_request():
#                 async with aiohttp.ClientSession() as session:
#                     async with session.post(create_kb_url, json=knowledge_base_data, headers=headers) as response:
#                         return await response.json(), response.status

#             # Make the async request synchronous
#             response_data, status_code = async_to_sync(async_request)()
#             return Response(response_data, status=status_code)
#         except Exception as e:
#             return Response({'error': f'Failed to create knowledge base: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)



# class SaveGoogleDriveCredentialsView(APIView):
#     permission_classes = [IsAuthenticated]

#     def post(self, request):
#         """
#         Endpoint to save Google Drive credentials for the authenticated user.
#         Expects the credentials (token, refresh_token, token_uri, client_id, client_secret, scopes) in the request payload.
#         """
#         # Extract the credentials from the request payload
#         credentials = request.data.get('credentials')
#         if not credentials:
#             return Response({'error': 'Missing credentials in request payload.'}, status=status.HTTP_400_BAD_REQUEST)

#         # Validate that all necessary fields are present in the credentials
#         required_fields = ['token', 'refresh_token', 'token_uri', 'client_id', 'client_secret', 'scopes']
#         if not all(field in credentials for field in required_fields):
#             return Response({'error': 'Invalid or incomplete credentials provided.'}, status=status.HTTP_400_BAD_REQUEST)

#         try:
#             # Save the credentials to the database associated with the user
#             GoogleDriveAccount.objects.update_or_create(
#                 user=request.user,
#                 defaults={'credentials': json.dumps(credentials)}
#             )

#             return Response({'message': 'Google Drive credentials saved successfully.'}, status=status.HTTP_200_OK)

#         except Exception as e:
#             # Handle any errors during the process
#             return Response({'error': f'An error occurred while saving credentials: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)





# Microsoft Sharepoint k endpoints


class MicrosoftInitView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        # Create a unique state parameter to prevent CSRF and store user info
        state = urlencode({'user_id': str(request.user.id), 'session_id': str(uuid.uuid4())})

        # Initialize the MSAL confidential client
        client = msal.ConfidentialClientApplication(
            client_id=settings.MICROSOFT_OAUTH_CLIENT_ID,
            client_credential=settings.MICROSOFT_OAUTH_CLIENT_SECRET,
            authority=f"https://login.microsoftonline.com/{settings.MICROSOFT_OAUTH_TENANT_ID}"
        )

        # Generate the authorization URL
        auth_url = client.get_authorization_request_url(
            scopes=settings.MICROSOFT_OAUTH_SCOPES,
            redirect_uri=settings.MICROSOFT_OAUTH_REDIRECT_URI,
            state=state,
            prompt='consent'  # Forces consent screen every time
        )

        # Return the auth URL for the frontend to redirect the user
        return Response({'auth_url': auth_url}, status=status.HTTP_200_OK)

from django.core.cache import cache

class MicrosoftCallbackView(APIView):
    permission_classes = [AllowAny]

    def get(self, request):
        # Extract code and state from query parameters
        code = request.GET.get('code')
        state = request.GET.get('state')
        print("c1")
        if not code or not state:
            return Response({'error': 'No code or state provided.'}, status=status.HTTP_400_BAD_REQUEST)

        # Parse state to extract user_id
        try:
            state_data = dict(pair.split('=') for pair in state.split('&'))
            user_id = state_data.get('user_id')
            session_id = state_data.get('session_id')

            # Check if session_id has been used
            if cache.get(session_id):
                return Response({'error': 'This session has already been processed.'}, status=status.HTTP_400_BAD_REQUEST)

            # Mark session_id as used for a limited time
            cache.set(session_id, 'used', timeout=300)  # Store for 5 minutes to prevent duplicates
        except Exception as e:
            return Response({'error': 'Invalid state parameter.'}, status=status.HTTP_400_BAD_REQUEST)

        # Retrieve the user
        User = get_user_model()
        try:
            user = User.objects.get(id=user_id)
        except User.DoesNotExist:
            return Response({'error': 'User not found.'}, status=status.HTTP_404_NOT_FOUND)

        # Initialize the MSAL confidential client
        client = msal.ConfidentialClientApplication(
            client_id=settings.MICROSOFT_OAUTH_CLIENT_ID,
            client_credential=settings.MICROSOFT_OAUTH_CLIENT_SECRET,
            authority=f"https://login.microsoftonline.com/{settings.MICROSOFT_OAUTH_TENANT_ID}"
        )

        try:
            # Acquire token by authorization code
            result = client.acquire_token_by_authorization_code(
                code=code,
                scopes=settings.MICROSOFT_OAUTH_SCOPES,
                redirect_uri=settings.MICROSOFT_OAUTH_REDIRECT_URI
            )

            if 'error' in result:
                return Response({'error': result.get('error_description')}, status=status.HTTP_400_BAD_REQUEST)

            # Extract tokens
            access_token = result.get('access_token')
            refresh_token = result.get('refresh_token')  # May be None depending on scopes and settings
            expires_in = result.get('expires_in')
            scope = result.get('scope')

            # Prepare credentials data
            creds_data = {
                'access_token': access_token,
                'refresh_token': refresh_token,
                'expires_in': expires_in,
                'scope': scope,
                'token_type': result.get('token_type'),
                'expires_at': result.get('expires_on')
            }

            # Store credentials in the database
            MicrosoftAccount.objects.update_or_create(
                user=user,
                defaults={'credentials': json.dumps(creds_data)}
            )

            # Generate JWT token for the user (if needed)
            jwt_payload = {
                'user_id': user.id
            }
            token = jwt.encode(jwt_payload, settings.SECRET_KEY, algorithm='HS256')

            # Redirect to the frontend with JWT and Microsoft access token in query parameters
            frontend_url = 'https://ai-rag-client-admin-git-staging-octalooptechnologies-projects.vercel.app/knowledgebase/own'
            params = {
                'token': token,
                'access_token': access_token,  # Microsoft access token
                # 'refresh_token': refresh_token  # Microsoft refresh token (optional)
            }
            redirect_url = f"{frontend_url}?{urlencode(params)}"

            print(creds_data)
            return HttpResponseRedirect(redirect_url)
            # return Response({
            #     'message': 'Microsoft sharepoint successfully authenticated and tokens saved.',
            #     'credentials': creds_data
            # }, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': 'An error occurred: ' + str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)







##chatbots k endpoints


class CreateChatbotView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        user_admin = request.user

        if not user_admin.is_user_admin:
            return Response({'error': 'Only User Admins can create chatbots'}, status=status.HTTP_403_FORBIDDEN)

        name = request.data.get('name')
        description = request.data.get('description', '')
        instructions = request.data.get('instructions', '')
        conversation_starter = request.data.get('conversation_starter', '')
        chatbot_profile_url = request.data.get('chatbot_profile_url', '')
        knowledge_base_id = request.data.get('knowledge_base_id')
        temperature = request.data.get('temperature', 0.7)
        max_tokens = request.data.get('max_tokens', 150)
        top_p = request.data.get('top_p', 0.9)
        category = request.data.get('category', '')
        model_name=request.data.get('model_name', 'gpt-3.5-turbo')

        if not name or not chatbot_profile_url or not knowledge_base_id:
            return Response({'error': 'Name, Profile URL, and Knowledge Base are required fields'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            knowledge_base = get_object_or_404(KnowledgeBase, id=knowledge_base_id, created_by=user_admin)

            chatbot = Chatbot.objects.create(
                name=name,
                model_name=model_name,
                description=description,
                instructions=instructions,
                conversation_starter=conversation_starter,
                chatbot_profile_url=chatbot_profile_url,
                created_by=user_admin,
                knowledge_base=knowledge_base,
                temperature=temperature,
                max_tokens=max_tokens,
                top_p=top_p,
                category=category
            )

            return Response({
                'message': 'Chatbot created successfully.',
                'chatbot_id': chatbot.id,
                'name': chatbot.name,
            
            }, status=status.HTTP_201_CREATED)
        
        except KnowledgeBase.DoesNotExist:
            return Response({'error': 'Knowledge Base not found or does not belong to you'}, status=status.HTTP_404_NOT_FOUND)

        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        



class UpdateChatbotView(APIView):
    permission_classes = [IsAuthenticated]

    def put(self, request, chatbot_id):
        user = request.user

        # Check if the user is an admin or has the 'editor' role
        if not (user.is_user_admin or user.role == 'editor'):
            return Response({'error': 'Only User Admins or Editors can update chatbots'}, status=status.HTTP_403_FORBIDDEN)

        # If the user is an editor, ensure they are associated with the User Admin who created the chatbot
        if user.role == 'editor':
            # Find the chatbot
            chatbot = get_object_or_404(Chatbot, id=chatbot_id)

            # # Check if the editor is associated with the user admin who created the chatbot
            # if not UserAdminUserRelationship.objects.filter(user=user, user_admin=chatbot.created_by).exists():
            #     return Response({'error': 'Editor is not associated with the User Admin who created this chatbot'}, status=status.HTTP_403_FORBIDDEN)
        else:
            # If the user is a User Admin, they can directly access the chatbot
            chatbot = get_object_or_404(Chatbot, id=chatbot_id, created_by=user)

        # Update chatbot fields with the data from the request
        chatbot.name = request.data.get('name', chatbot.name)
        chatbot.model_name=request.data.get('model_name', chatbot.model_name)
        chatbot.description = request.data.get('description', chatbot.description)
        chatbot.instructions = request.data.get('instructions', chatbot.instructions)
        chatbot.conversation_starter = request.data.get('conversation_starter', chatbot.conversation_starter)
        chatbot.chatbot_profile_url = request.data.get('chatbot_profile_url', chatbot.chatbot_profile_url)
        chatbot.temperature = request.data.get('temperature', chatbot.temperature)
        chatbot.max_tokens = request.data.get('max_tokens', chatbot.max_tokens)
        chatbot.top_p = request.data.get('top_p', chatbot.top_p)
        chatbot.category = request.data.get('category', chatbot.category)

        # Check and update the knowledge base if provided
        knowledge_base_id = request.data.get('knowledge_base_id')
        if knowledge_base_id:
            knowledge_base = get_object_or_404(KnowledgeBase, id=knowledge_base_id, created_by=chatbot.created_by)
            chatbot.knowledge_base = knowledge_base

        # Save the chatbot changes
        chatbot.save()

        return Response({
            'message': 'Chatbot updated successfully.',
            'chatbot_id': chatbot.id,
            'name': chatbot.name
        }, status=status.HTTP_200_OK)




class DeleteChatbotView(APIView):
    permission_classes = [IsAuthenticated]

    def delete(self, request, chatbot_id):
        user_admin = request.user

        if not user_admin.is_user_admin:
            return Response({'error': 'Only User Admins can delete chatbots'}, status=status.HTTP_403_FORBIDDEN)

        chatbot = get_object_or_404(Chatbot, id=chatbot_id, created_by=user_admin)
        chatbot.delete()

        return Response({'message': 'Chatbot deleted successfully.'}, status=status.HTTP_200_OK)
    


class ListChatbotsView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        user_admin = request.user

        if not user_admin.is_user_admin:
            return Response({'error': 'Only User Admins can view their chatbots'}, status=status.HTTP_403_FORBIDDEN)

        chatbots = Chatbot.objects.filter(created_by=user_admin)
        chatbots_data = [{
            'id': chatbot.id,
            'model_name':chatbot.model_name,
            'name': chatbot.name,
            'description': chatbot.description,
            'instructions': chatbot.instructions,
            'conversation_starter': chatbot.conversation_starter,
            'chatbot_profile_url': chatbot.chatbot_profile_url,
            'temperature': chatbot.temperature,
            'max_tokens': chatbot.max_tokens,
            'top_p': chatbot.top_p,
            'created_at': chatbot.created_at,
            'category': chatbot.category,
            'created_by':{
                'id':chatbot.created_by.id,
                'username':chatbot.created_by.username,
                'email': chatbot.created_by.email,

            }
        } for chatbot in chatbots]

        return Response({'chatbots': chatbots_data}, status=status.HTTP_200_OK)
    



class GetChatbotByIdView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, chatbot_id):
        user = request.user

        # Check if the user is an admin or has the 'editor' role
        if not (user.is_user_admin or user.role == 'editor'):
            return Response({'error': 'Only User Admins or Editors can view this chatbot'}, status=status.HTTP_403_FORBIDDEN)

        # If the user is an editor, ensure they are associated with the User Admin who created the chatbot
        if user.role == 'editor':
            # Find the chatbot first
            chatbot = get_object_or_404(Chatbot, id=chatbot_id)

            # Check if the editor is associated with the user admin who created the chatbot
            # if not UserAdminUserRelationship.objects.filter(user=user, user_admin=chatbot.created_by).exists():
            #     return Response({'error': 'Editor is not associated with the User Admin who created this chatbot'}, status=status.HTTP_403_FORBIDDEN)
        else:
            # If the user is a User Admin, they can directly access the chatbot
            chatbot = get_object_or_404(Chatbot, id=chatbot_id, created_by=user)

        # Prepare chatbot data
        chatbot_data = {
            'id': chatbot.id,
            'name': chatbot.name,
            'model_name':chatbot.model_name,
            'description': chatbot.description,
            'instructions': chatbot.instructions,
            'conversation_starter': chatbot.conversation_starter,
            'chatbot_profile_url': chatbot.chatbot_profile_url,
            'knowledge_base_id': chatbot.knowledge_base.id if chatbot.knowledge_base else None,  # Return the ID of the knowledge base
            'temperature': chatbot.temperature,
            'max_tokens': chatbot.max_tokens,
            'top_p': chatbot.top_p,
            'created_at': chatbot.created_at,
            'category': chatbot.category,
            'created_by': {
                'id': chatbot.created_by.id,
                'username': chatbot.created_by.username,
                'email': chatbot.created_by.email
            }
        }

        return Response({'chatbot': chatbot_data}, status=status.HTTP_200_OK)


        







##update user roles:
class UpdateUserRoleView(APIView):
    permission_classes = [IsAuthenticated]

    def put(self, request, user_id):
        user_admin = request.user

        # Ensure the current user is an admin
        if not user_admin.is_user_admin:
            return Response({'error': 'Only User Admins can update user roles.'}, status=status.HTTP_403_FORBIDDEN)

        # Fetch the user to update their role
        user = get_object_or_404(User, id=user_id)

        # Ensure the user to be updated is associated with the current admin
        if not UserAdminUserRelationship.objects.filter(user=user, user_admin=user_admin).exists():
            return Response({'error': 'You do not have permission to update this user\'s role.'}, status=status.HTTP_403_FORBIDDEN)

        # Get the new role from the request
        new_role = request.data.get('role')

        # Ensure the role is valid
        if new_role not in dict(User.ROLE_CHOICES):
            return Response({'error': 'Invalid role specified.'}, status=status.HTTP_400_BAD_REQUEST)

        # Update the user's role
        user.role = new_role
        user.save()

        return Response({
            'message': f"User role updated to {new_role}.",
            'user_id': user.id,
            'new_role': user.role
        }, status=status.HTTP_200_OK)
    





class GetUserDetailsView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, user_id):
        # Fetch the user by ID
        user = get_object_or_404(User, id=user_id)

        # Fetch the user's role
        user_role = user.role

        # Get the user admin details using the UserAdminUserRelationship
        user_admin_relationship = UserAdminUserRelationship.objects.filter(user=user).first()
        
        if not user_admin_relationship:
            return Response({
                'error': 'This user does not have an associated admin.'
            }, status=status.HTTP_400_BAD_REQUEST)

        # Get the admin details
        admin = user_admin_relationship.user_admin

        # Fetch knowledge bases created by this user
        knowledge_bases = KnowledgeBase.objects.filter(created_by=user)
        
        # Serialize knowledge bases to get their IDs and names
        knowledge_base_list = [
            {'knowledgebase_id': kb.id, 'knowledgebase_name': kb.name} for kb in knowledge_bases
        ]

        # Count the number of knowledge bases created by this user
        knowledge_base_count = knowledge_bases.count()

        # Prepare the response data
        response_data = {
            'user_id': user.id,
            'username': user.username,
            'email': user.email,
            "profile_photo_url": user.profile_photo_url,
            'role': user_role,
            'admin_details': {
                'admin_id': admin.id,
                'admin_username': admin.username,
                'admin_email': admin.email
            },
            'knowledge_base_count': knowledge_base_count,
            'knowledge_bases': knowledge_base_list
        }

        return Response(response_data, status=status.HTTP_200_OK)
    





###palying with pinecone vector


from pinecone import Pinecone

class ViewNamespaceView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        logger.info("Incoming request to fetch namespace vectors with pagination")

        # Extract request data
        admin_id = request.data.get('admin_id')
        knowledge_base_id = request.data.get('knowledge_base_id')
        namespace = request.data.get('namespace')
        pagination_token = request.data.get('pagination_token', None)  # Get pagination token if provided
        limit = request.data.get('limit', 10)  # Get limit, defaulting to 10 if not provided

        # Validate request data
        if not all([admin_id, knowledge_base_id, namespace]):
            return Response({'error': 'Missing required fields: admin_id, knowledge_base_id, and namespace'}, status=status.HTTP_400_BAD_REQUEST)

        # Retrieve the knowledge base from the database
        try:
            knowledge_base = KnowledgeBase.objects.get(id=knowledge_base_id, namespace=namespace)
        except KnowledgeBase.DoesNotExist:
            return Response({'error': 'Knowledge base not found'}, status=status.HTTP_404_NOT_FOUND)

        # Check permissions and get user admin
        user = request.user
        if user.is_user_admin:
            user_admin = user
        else:
            # Ensure the user is an editor or associated with the user admin
            if user.role != 'editor':
                return Response({'error': 'Only User Admins or Editors can access vectors'}, status=status.HTTP_403_FORBIDDEN)

            try:
                user_admin = User.objects.get(id=admin_id, is_user_admin=True)
            except User.DoesNotExist:
                return Response({'error': 'Admin not found'}, status=status.HTTP_404_NOT_FOUND)

            # Check if the editor is associated with the admin
            if not UserAdminUserRelationship.objects.filter(user=user, user_admin=user_admin).exists():
                return Response({'error': 'You are not associated with the provided admin.'}, status=status.HTTP_403_FORBIDDEN)

        # Check if the admin has a Pinecone index
        index_name = user_admin.pinecone_index
        if not index_name:
            return Response({'error': 'Admin does not have a Pinecone index'}, status=status.HTTP_400_BAD_REQUEST)

        # Initialize Pinecone and fetch vectors with pagination and limit
        try:
            api_settings = APISettings.objects.first()
            if not api_settings:
                        raise ImproperlyConfigured("APISettings instance not found.")
                # Initialize Pinecone and clear the existing namespace
            OPENAI_API_KEY = api_settings.openai_api_key or os.getenv("OPENAI_API_KEY")
            PINECONE_API_KEY = api_settings.pinecone_api_key or os.getenv("PINECONE_API_KEY")
            pinecone_initializer = PineconeInitializer(pinecone_api=PINECONE_API_KEY, open_ai_api=OPENAI_API_KEY)
            logger.info(f"Fetching vectors for namespace: {namespace} with pagination token: {pagination_token} and limit: {limit}")

            # Pass the pagination token and limit to fetch the vectors
            vectors_with_metadata = pinecone_initializer.list_vectors_with_metadata(
                index_name=index_name,
                namespace=namespace,
                prefix='',
                pagination_token=pagination_token,  # Use pagination token if provided
                limit=limit  # Limit the number of vectors per page
            )

            # Extract relevant data from the response
            namespace_done = vectors_with_metadata.get('namespace')
            vectors = vectors_with_metadata.get("vectors")
            next_page_token = vectors_with_metadata.get('pagination', {}).get('next_page_token')

        except Exception as e:
            logger.error(f"Error fetching vectors from Pinecone: {str(e)}", exc_info=True)
            return Response({'error': f'Error fetching vectors: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        # Return the structured response with pagination token if available
        if vectors:
            return Response({
                'namespace': namespace_done,
                'vectors': vectors,
                'next_page_token': next_page_token  # Return the next page token if more data exists
            }, status=status.HTTP_200_OK)
        else:
            return Response({'error': 'No vectors found in the specified namespace.'}, status=status.HTTP_404_NOT_FOUND)



class UpdateVector(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        # Retrieve required data from request data (POST body)
        admin_id = request.data.get('admin_id')
        knowledge_base_id = request.data.get('knowledge_base_id')
        namespace = request.data.get('namespace')
        vector_id = request.data.get('vector_id')
        updated_text = request.data.get('updated_text')

        # Validate request data
        if not admin_id or not knowledge_base_id or not namespace or not vector_id or not updated_text:
            return Response({'error': 'Missing required fields: admin_id, knowledge_base_id, namespace, vector_id, and updated_text'}, status=status.HTTP_400_BAD_REQUEST)

        # Check if the knowledge base exists
        try:
            knowledge_base = KnowledgeBase.objects.get(id=knowledge_base_id, namespace=namespace)
        except KnowledgeBase.DoesNotExist:
            return Response({'error': 'Knowledge base not found'}, status=status.HTTP_404_NOT_FOUND)

        # Check permissions and get user admin
        user = request.user
        if user.is_user_admin:
            user_admin = user
        else:
            if user.role != 'editor':
                return Response({'error': 'Only User Admins or Editors can update vectors'}, status=status.HTTP_403_FORBIDDEN)

            # Check if the admin_id exists and is valid
            try:
                user_admin = User.objects.get(id=admin_id, is_user_admin=True)
            except User.DoesNotExist:
                return Response({'error': 'Admin not found'}, status=status.HTTP_404_NOT_FOUND)

            # Ensure the editor is associated with the admin
            if not UserAdminUserRelationship.objects.filter(user=user, user_admin=user_admin).exists():
                return Response({'error': 'You are not associated with the provided admin.'}, status=status.HTTP_403_FORBIDDEN)

        # Check if the admin has a Pinecone index
        index_name = user_admin.pinecone_index
        if not index_name:
            return Response({'error': 'Admin does not have a Pinecone index'}, status=status.HTTP_400_BAD_REQUEST)

        # Initialize Pinecone and update vector
        try:
            api_settings = APISettings.objects.first()
            if not api_settings:
                        raise ImproperlyConfigured("APISettings instance not found.")
                # Initialize Pinecone and clear the existing namespace
            OPENAI_API_KEY = api_settings.openai_api_key or os.getenv("OPENAI_API_KEY")
            PINECONE_API_KEY = api_settings.pinecone_api_key or os.getenv("PINECONE_API_KEY")
            pinecone_initializer = PineconeInitializer(pinecone_api=PINECONE_API_KEY, open_ai_api=OPENAI_API_KEY)
            updated_vector = pinecone_initializer.update_vector_data(
                environment="us-central1-gcp",
                index_name=index_name,
                vector_id=vector_id,
                updated_text=updated_text,
                namespace=namespace
            )
            print(updated_vector)
        except Exception as e:
            return Response({'error': f'Error updating vector: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        # Return the updated vector metadata
        return Response({
            'message': 'Vector updated successfully.',
        }, status=status.HTTP_200_OK)
    




class DeleteVectorView(APIView):
    permission_classes = [IsAuthenticated]  # Adjust permissions as needed

    def delete(self, request, admin_id, vector_id, namespace):
        # Log the incoming delete request
        logger.info(f"Delete request for vector_id: {vector_id}, index: {admin_id}, namespace: {namespace}")
        try:
            user_admin = User.objects.get(id=admin_id, is_user_admin=True)
        except Exception as e:
            return Response({'error': 'no access to this knowledgebase . either admin is delete or knowledgebase'}, status=status.HTTP_404_NOT_FOUND)
        # Get the user (only user admins or editors should be able to delete vectors)
        index_name=user_admin.pinecone_index
        user = request.user
        if not user.is_user_admin and user.role != 'editor':
            logger.error(f"Unauthorized attempt to delete vector by user {user.id}")
            return Response({'error': 'Only User Admins or Editors can delete vectors'}, status=status.HTTP_403_FORBIDDEN)

        # Initialize Pinecone and attempt to delete the vector
        try:
            index_name=user_admin.pinecone_index
            api_settings = APISettings.objects.first()
            if not api_settings:
                        raise ImproperlyConfigured("APISettings instance not found.")
                # Initialize Pinecone and clear the existing namespace
            OPENAI_API_KEY = api_settings.openai_api_key or os.getenv("OPENAI_API_KEY")
            PINECONE_API_KEY = api_settings.pinecone_api_key or os.getenv("PINECONE_API_KEY")
            pinecone_initializer = PineconeInitializer(pinecone_api=PINECONE_API_KEY, open_ai_api=OPENAI_API_KEY)
            pinecone_initializer.delete_vector_from_pinecone(index_name=index_name, vector_id=vector_id, namespace=namespace)
            logger.info(f"Vector {vector_id} successfully deleted from namespace {namespace} in index {index_name}")
            return Response({'message': f'Vector {vector_id} successfully deleted from namespace {namespace} in index {index_name}'}, status=status.HTTP_200_OK)
        
        except Exception as e:
            logger.error(f"Error deleting vector {vector_id}: {str(e)}")
            return Response({'error': f'Error deleting vector: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
