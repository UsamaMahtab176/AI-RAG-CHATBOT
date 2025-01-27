from dotenv import load_dotenv
import os
from pinecone import Pinecone, ServerlessSpec, PodSpec
import time
from langchain_openai import OpenAIEmbeddings
from langchain_pinecone import PineconeVectorStore
from langchain_text_splitters import CharacterTextSplitter
from django.conf import settings
from langchain_core.documents import Document
from typing import List
import base64
import io
from PyPDF2 import PdfReader
from docx import Document
import pdfplumber
import logging
logger = logging.getLogger(__name__)

#teting
class PineconeInitializer:
    def __init__(self,pinecone_api,open_ai_api):
        self.pinecone_api_key = pinecone_api
        self.use_serverless = True
        self.openai_key=open_ai_api
        try:
            self.pc = Pinecone(api_key=pinecone_api)
        except ValueError as e:
            print(f"Failed to initialize Pinecone client: {e}")
            raise ValueError("Pinecone API key is missing or invalid")



    def initialize_pinecone(self, index_name):
        

        if self.use_serverless:
            spec = ServerlessSpec(cloud='aws', region='us-east-1')
        else:
            spec = PodSpec(
                environment="us-central1-gcp",
                pod_type="p1.x1",
                pods=1
            )

        if index_name in self.pc.list_indexes().names():
            # self.pc.delete_index(index_name)
            print ("index already exist")
        else :    
            self.pc.create_index(
                index_name,
                dimension=1536,  # dimensionality of text-embedding-ada-002
                metric='dotproduct',
                spec=spec
            )

        while not self.pc.describe_index(index_name).status['ready']:
            time.sleep(1)

        return self.pc.Index(index_name)
    

    def namespace_exists_in_index(self, index_name, namespace):
        """
        Check if a namespace exists in the specified Pinecone index.
        Returns True if the namespace exists, otherwise False.
        """
        try:
            index = self.pc.Index(index_name)  # Connect to the specified index
            # Query the index with a small vector and check if the namespace returns results
            response = index.query(
                vector=[0.0] * 512,  # Example dummy vector, adjust if necessary for your vector dimensions
                top_k=1,
                namespace=namespace
            )
            # If results are returned, the namespace exists
            return len(response['matches']) > 0
        except Exception as e:
            print(f"Error while checking namespace existence: {str(e)}")
            return False
        
    def show_indexes (self):
        print (self.pc.list_indexes())
        new =self.pc.list_indexes()
        return new

    def describe_index(self,index_name):
        print (self.pc.describe_index(index_name))
        return self.pc.describe_index(index_name)
        
        
    def delete_index_pinecone(self, index_name):
        self.pc.delete_index(index_name)



    def Embeding_Pdf_to_pincecone(self,loader,index_name):
        data = loader
        print("now data is   :  ",data)
        text_splitter = CharacterTextSplitter(chunk_size=50, chunk_overlap=4)
        docs = text_splitter.split_documents(data)
        model_name = 'text-embedding-ada-002'
        embeddings = OpenAIEmbeddings(
        model=model_name,
        openai_api_key=self.openai_key
        )
        # index_name = "langchain-retrieval-augmentation-fast"
        vectorstore = PineconeVectorStore.from_documents(docs, embeddings, index_name=index_name)
        print ("Embdeding Done ")
        return vectorstore
    
    def Embeding_Text_list_to_pinecone(self, texts, index_name, Agent_id, namespace, metadata_list=None):
        try:
            model_name = 'text-embedding-ada-002'
            embeddings = OpenAIEmbeddings(
                model=model_name,
                openai_api_key=self.openai_key
            )

            # Ensure metadata_list has an entry for each text chunk
            if metadata_list is None:
                metadata_list = [{"Admin_id": Agent_id} for _ in texts]
            else:
                # Add the Agent_id to each metadata dictionary if not already included
                for metadata in metadata_list:
                    if "Admin_id" not in metadata:
                        metadata["Admin_id"] = Agent_id

            # Embed the chunks with their respective metadata
            vectorstore_from_texts = PineconeVectorStore.from_texts(
                texts=texts,
                index_name=index_name,
                embedding=embeddings,
                metadatas=metadata_list,  # Pass the metadata here
                namespace=namespace
            )

            # Logging instead of print for production
            logger.info("Embedding completed successfully for index: %s", index_name)
            return vectorstore_from_texts

        except Exception as e:
            # Log the exception with traceback for better diagnostics
            logger.exception("Error embedding text list to Pinecone: %s", str(e))
            return None
           
    def texts_to_documents(texts: List[str], source: str) -> List[Document]:
        documents = []
        for i, text in enumerate(texts):
            doc = Document(
                page_content=text,
                metadata={
                    'source': source,
                    'page': i
                }
            )
            documents.append(doc)
        return documents
        
    def connect_to_index(self,index_name):
        index = self.pc.Index(index_name)
        time.sleep(1)
        return index
    


    def extract_and_concatenate_text(documents):
        """
        Extract text from each document and concatenate them with a logical separator.
        """
        all_text = ""  # Initialize an empty string to hold all text
        
        for document_data in documents:
            base64_file_content = document_data.get('file')  # base64 string
            document_type = document_data.get('document_type').lower()

            if not document_type or not base64_file_content:
                raise ValueError(f'Missing file or document_type')

            # Decode base64 content to bytes
            if base64_file_content.startswith('data:'):
                base64_file_content = base64_file_content.split(',')[1]
            try:
                file_content = base64.b64decode(base64_file_content)
            except Exception as e:
                raise ValueError(f'Invalid base64 file content for document: {e}')

            # Process document content based on type
            if document_type == 'pdf':
                pdf_reader = PdfReader(io.BytesIO(file_content))
                text = ""
                for page_num in range(len(pdf_reader.pages)):
                    text += pdf_reader.pages[page_num].extract_text()
                all_text += text + "\n\n"  # Add separator between documents

            elif document_type == 'docx':
                doc_stream = io.BytesIO(file_content)
                doc = Document(doc_stream)
                text = "\n".join([para.text for para in doc.paragraphs])
                all_text += text + "\n\n"  # Add separator between documents

            elif document_type == 'text':
                text = file_content.decode('utf-8')
                all_text += text + "\n\n"  # Add separator between documents

        return all_text.strip()  # Remove trailing newlines




    def universal_text_splitter(text, chunk_size=50, chunk_overlap=10):
        """
        Split text using a variety of common separators (paragraphs, sentences, etc.).
        :param text: The raw text content to split.
        :param chunk_size: The max size of each chunk.
        :param chunk_overlap: The number of characters that overlap between chunks.
        :return: A list of split text chunks.
        """
        # Split by paragraph (newlines)
        paragraphs = text.split("\n\n")  # Splitting on double newlines
        
        chunks = []
        current_chunk = ""

        for paragraph in paragraphs:
            if len(paragraph) + len(current_chunk) > chunk_size:
                # Save the current chunk when exceeding the chunk size
                chunks.append(current_chunk.strip())

                # Start a new chunk with an overlap
                current_chunk = paragraph[:chunk_overlap]

            # Add paragraph to the current chunk
            current_chunk += paragraph
        
        # Add any remaining text as the last chunk
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks


    def delete_namespace_from_pinecone(self, index_name, namespace):
        # Initialize the Pinecone index using gRPC
        index = self.pc.Index(index_name)

        try:
            # Delete all vectors from the specified namespace
            # If namespace exists, this deletes all vectors within the namespace
            index.delete(delete_all=True, namespace=namespace)
        except Exception as e:
            raise Exception(f"Error deleting namespace '{namespace}' from index '{index_name}': {str(e)}")
        


    def list_vectors_with_metadata(self, index_name, namespace, limit, pagination_token,prefix=''):
    # Initialize Pinecone client
        pc = self.pc
        # Connect to the specified index
        index = pc.Index(index_name)

        # Fetch the first page of vectors with metadata
        results = index.list_paginated(
            prefix=prefix,   # Optional prefix filter for vector IDs
            limit=limit,       # Limit of vectors to retrieve in one go (adjust as needed)
            namespace=namespace,
            pagination_token=pagination_token
        )

        # Prepare response dictionary
        response_data = {
            "namespace": namespace,
            "vectors": [],
            "pagination": None,
            "api_usage": None
        }

        # Check if vectors exist and iterate through each vector
        if results.vectors:
            for vector in results.vectors:
                # Fetch vector metadata using the fetch method
                fetched_vector = index.fetch([vector.id], namespace=namespace)
                if vector.id in fetched_vector.vectors:
                    vector_data = fetched_vector.vectors[vector.id]
                    metadata = vector_data.metadata

                    # Ensure the metadata is serializable (handle non-serializable types if needed)
                    # Append vector info to the response
                    response_data["vectors"].append({
                        "vector_id": vector.id,
                        "metadata": metadata  # Ensure metadata is a dict or serializable type
                    })

        # Handle pagination if available
        if results.pagination and results.pagination.next:
            response_data["pagination"] = {
                "next_page_token": results.pagination.next
            }

        # Handle API usage information
        if results.usage:
            response_data["api_usage"] = results.usage

        return response_data


    def update_vector_data(self, environment, index_name, vector_id, updated_text, namespace):
            # Initialize Pinecone client
            pc = Pinecone(api_key=self.pinecone_api_key, environment=environment)
            
            # Connect to your Pinecone index
            index = pc.Index(index_name)
            
            # Fetch the existing vector to verify its existence within the namespace
            existing_vector = index.fetch(ids=[vector_id], namespace=namespace)
            print("Existing vector metadata: ", existing_vector)
            
            # New metadata with the updated text
            new_metadata = {
                "page": 0,
                "source": index_name,  # Assuming source is the same as the index name
                "text": updated_text
            }
            
            # Update the vector's metadata in the specified namespace
            index.update(id=vector_id, set_metadata=new_metadata, namespace=namespace)
            
            # Fetch the updated vector to confirm changes within the namespace
            updated_vector = index.fetch(ids=[vector_id], namespace=namespace)
            print("Updated vector metadata: ", updated_vector)
            
            # Return updated vector metadata
            return updated_vector
    


    def delete_vector_from_pinecone(self, index_name, vector_id, namespace):
    # Initialize Pinecone client
        pc = self.pc
        
        # Connect to the specified index
        index = pc.Index(index_name)
        
        # Delete the vector by its ID and namespace
        try:
            index.delete(ids=[vector_id], namespace=namespace)
            print(f"Vector {vector_id} successfully deleted from namespace '{namespace}' in index '{index_name}'.")
        except Exception as e:
            print(f"Error deleting vector {vector_id}: {str(e)}")





 

# def main():
#     load_dotenv(override=True)
#     pinecone_api_key = os.environ.get('PINECONE_API_KEY')
#     pine_cone_initializer = PineconeInitializer(pinecone_api_key)
#     pine_cone_initializer.show_indexes(index_name="langchain-retrieval-augmentation-fast")
#     # openai_api_key = os.environ.get('OPENAI_API_KEY')
#     # if not pinecone_api_key or not openai_api_key:
#     #     raise ValueError("Pinecone API key or OpenAI API key not found in environment variables.")
#     # # Initialize Pinecone
#     # pinecone_initializer = PineconeInitializer(pinecone_api_key)
#     # index_name = 'langchain-retrieval-augmentation-fast'
#     # print ("Creating index .....")
#     # index = pinecone_initializer.initialize_pinecone(index_name) 
#     # print ("Index has been created....")

# if __name__ == "__main__":
#     main()